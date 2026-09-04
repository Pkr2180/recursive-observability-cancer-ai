from __future__ import annotations

import csv
import math
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "manuscript_package"
QA = OUT / "qa"
ASSETS = ROOT / "manuscript_assets"
WORKFLOW_FIG = ASSETS / "figure_01_recursive_observability_workflow.png"
RESULTS = ROOT / "retrieved_results_2026-08-25" / "results"
PUB = RESULTS / "observability_primary" / "publication_package"
PRIMARY_FIG = PUB / "figures"
PRIMARY_TAB = PUB / "tables"
REINF = RESULTS / "reinforcement_master_brain"
REINF_FIG = REINF / "figures"

MAIN_DOCX = OUT / "Nature_Style_Master_Reinforcement_AI_Cancer_Futures_Manuscript.docx"
SUPP_DOCX = OUT / "Nature_Style_Master_Reinforcement_AI_Cancer_Futures_Supplement.docx"

NAVY = "17365D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "000000"


TITLE = (
    "A Recursively Observable Master Reinforcement AI Reveals "
    "Unstable Cancer Futures"
)
SHORT_TITLE = "Recursive observability of unstable cancer futures"


ABSTRACT = (
    "Cancer is commonly represented as a molecular state to be classified, although malignant systems are "
    "history-dependent populations that branch, compensate and reorganize after perturbation. Artificial "
    "intelligence compounds this reduction by compressing discordant molecular views into one prediction and "
    "treating internal instability as disposable error. Here we introduce a different scientific object: a "
    "recursively observable Master Reinforcement AI that uses the instability of its own distributed observer "
    "states as a testable sensing surface for unstable cancer behaviour. Five synchronized streams—observed "
    "cancer-state change, individual-observer instability, inter-observer disagreement, Master instability and "
    "uncertainty in the Master's self-observation—were preserved under a fail-closed 20-field transition contract. "
    "Across DepMap–PRISM molecular/pharmacological states, LINCS matched-time perturbations and a TCGA held-out "
    "cancer-project outcome audit, all 4,431 real-data transitions were trace-complete. Observer–system instability "
    "coupling was 0.271, 0.895 and 0.896, respectively. In LINCS, biological change and Master error remained "
    "coupled under linear, rank and nonlinear observed-value null tests (empirical P=0.001; false-discovery rate "
    "0.001), while low state reconstructability (0.223) exposed where the forecaster failed. In TCGA, moderate "
    "held-out discrimination (AUROC 0.677) coexisted with failed self-error ranking (r=-0.267), demonstrating that "
    "performance can appear acceptable while a model cannot identify its unsafe predictions. Together, these "
    "results establish a statistically resolved and previously unmeasured failure phenotype: biological-state "
    "instability, forecast error and the reliability of self-observation can diverge and must be evaluated as "
    "separate scientific endpoints. An offline "
    "reinforcement extension learned to reweight observers and defer, but did not decisively improve reward over "
    "the frozen Master. The advance is therefore not a claim of clinical autonomy or a marginal score increase. "
    "It is a new epistemic and "
    "engineering standard: cancer AI should reveal which futures remain open, why its internal observers diverge, "
    "how reliably it can observe that divergence and when instability requires falsification, more measurement or "
    "abstention. The framework also defines an unknown-feature boundary: unexplained latent structure is retained as a "
    "candidate for held-out and orthogonal validation, not prematurely labelled as novel cancer biology."
)


SECTIONS = [
    (
        "1. Introduction",
        [
            (
                "Cancer is not a fixed molecular object. It is an evolving population of states shaped by genetic "
                "variation, epigenetic plasticity, lineage constraints, microenvironmental pressures and treatment. "
                "Pan-cancer analyses have revealed strong cell-of-origin structure across thousands of tumours, yet "
                "they also expose recurring programmes that cross anatomical categories [1,2]. Within an individual "
                "tumour, clonal competition and reversible phenotypic transitions create alternative routes to "
                "survival [3-5]. Drug-tolerant persister states demonstrate that a population can temporarily occupy "
                "a protected phenotype without requiring a stable resistance mutation [6]. Consequently, a sample "
                "taken at one time is not the disease itself; it is one observation of a system that retains several "
                "possible futures."
            ),
            (
                "Modern cancer resources capture complementary slices of this system. TCGA establishes tumour-scale "
                "molecular and clinical variation [1,7]; DepMap measures context-specific genetic dependencies [8]; "
                "PRISM and GDSC map pharmacological response across cancer cell models [9,10]; and LINCS/L1000 records "
                "transcriptional responses to perturbation at scale [11]. These resources have enabled powerful "
                "classification and prediction, but their usual integration remains present-state centred. A model is "
                "asked to estimate a label, dependency, response or endpoint. When component models disagree, their "
                "predictions are often averaged, calibrated or discarded. The disagreement itself is rarely retained "
                "as a biological measurement."
            ),
            (
                "This creates a critical knowledge gap. A high-confidence prediction may conceal that different "
                "molecular views imply incompatible futures, while a low-confidence prediction may arise either from "
                "weak modelling or from a genuinely unstable biological state. Conventional performance metrics cannot "
                "separate these explanations because they evaluate the output without observing the internal process "
                "that produced it. Explainability after the event is insufficient: it describes a decision but does not "
                "establish whether agent states, transitions, counterfactual influence, uncertainty and self-error were "
                "visible while the decision evolved."
            ),
            (
                "We therefore tested a different proposition: structured instability inside a multi-agent AI system "
                "can become a computational sensor of unstable cancer behaviour if, and only if, every layer of that "
                "instability is separately observable and evaluated against null worlds. The framework does not equate "
                "AI uncertainty with biology. It measures five synchronized channels and asks whether their coupling "
                "exceeds what is expected after observed values are rearranged. The Master is recursively observed: the "
                "system records not only its state and error, but also uncertainty in its estimate of its own error."
            ),
            (
                "The study is different in four respects. First, scientific observability, not predictive accuracy, is "
                "the primary endpoint. Second, all biological execution uses observed public pan-cancer data; "
                "permutations form null distributions but do not simulate patients or trajectories. Third, whole "
                "perturbations and whole cancer projects are held out, reducing leakage across related observations. "
                "Fourth, failures remain visible. Low LINCS future-state reconstructability and failed TCGA self-error "
                "generalization are reported as scientific results rather than hidden behind an ensemble average. We "
                "hypothesized that this architecture would be trace-complete and falsifiable across molecular, dynamic "
                "and patient-outcome settings, and that cancer-AI instability coupling would identify states in which "
                "the model should adapt or defer."
            ),
            (
                "The position advanced here is therefore stronger than an ensemble-method proposal and narrower than "
                "a claim that AI has solved cancer. Cancer should be treated as a constrained field of possible futures, "
                "not only as a present molecular snapshot; a Master AI should be treated as a state of observer states, "
                "not only as an output layer; and instability should become a primary measurement, not a residual to be "
                "averaged away. The study tests the enabling requirement for that position: whether biological, agent, "
                "Master and self-observer dynamics can remain separate, reconstructable and falsifiable on real data. "
                "If they can, forecast failure itself can become scientific evidence about where the representation is "
                "under strain. If they cannot, the future-space theory fails at the level of its instrument. The "
                "complete recursive-observability workflow is summarized in Fig. 1."
            ),
        ],
    ),
    (
        "2. Materials and Methods",
        [],
    ),
    (
        "2.1 Study design and data policy",
        [
            (
                "We conducted an observational computational study using only public pan-cancer measurements. The "
                "analysis integrated cross-sectional molecular and pharmacological data, matched-time perturbational "
                "transcriptomics and baseline tumour transcriptomes linked to a fixed-horizon survival endpoint. No "
                "synthetic patient, cancer state or biological trajectory was admitted as evidence. Randomization was "
                "restricted to bootstrap resampling of observed biological units, sign-flipping or permutation of "
                "observed values for prespecified null tests. All execution was provenance-linked and versioned in a "
                "single remote workspace."
            ),
        ],
    ),
    (
        "2.2 Pan-cancer resources and harmonisation",
        [
            (
                "The tumour layer comprised 330 open-access TCGA STAR-count RNA-sequencing files balanced at ten files "
                "for each of 33 projects. Harmonisation yielded count and TPM matrices across 60,660 genes. The official "
                "TCGA Pan-Cancer Clinical Data Resource contained 11,160 patient records and linked to all 330 RNA files "
                "[7]. A censoring-aware two-year overall-survival endpoint retained 227 unique baseline-tumour patients "
                "across all projects. Entire cancer projects, rather than individual samples, were assigned to fit, "
                "calibration and test partitions; the untouched test set contained 42 patients from seven projects."
            ),
            (
                "The cell-model layer used DepMap 24Q2 expression and CRISPR dependency matrices and PRISM Repurposing "
                "Public 24Q2 drug response. Real-data validation covered 1,066 DepMap models, 20 sufficiently represented "
                "lineages, 100 candidate-gene observed-value permutation tests, 615 drug-target pairs and 1,033 "
                "lineage-specific dependency-pharmacology associations. The first shared Master analysis contained 723 "
                "cell models from 29 lineage labels. External replication used official GDSC1/GDSC2 fitted responses: "
                "575,197 measurements, including 322,040 rows mapped by exact Sanger identifiers to 531 DepMap models "
                "[10]."
            ),
            (
                "The dynamic layer used the checksum-verified NCBI GEO GSE70138 LINCS Level-5 matrix [11]. The source "
                "contained 118,050 signatures by 12,328 genes; cancer filtering retained 89,192 signatures and all 978 "
                "directly measured landmark genes. Exact matching by cell line, perturbation, dose and perturbation type "
                "produced 3,675 observed transitions across 3, 6 and 24 hours. Twenty-four complete perturbation "
                "identifiers were frozen before model evaluation, producing 712 unseen-perturbation test transitions."
            ),
        ],
    ),
    (
        "2.3 Multi-agent Master architectures",
        [
            (
                "Each architecture presented the same biological unit to three observers with different inductive or "
                "modality-specific views. The DepMap-PRISM Master integrated expression, CRISPR dependency and "
                "pharmacology agents. The LINCS Master integrated ridge, random-forest and nearest-neighbour forecasts "
                "of the later transcriptional state. The TCGA Master integrated linear-transcriptome, nonlinear-forest "
                "and patient-neighbourhood estimators. Agents were not forced to agree: their individual instability, "
                "pairwise divergence and counterfactual influence on the Master were retained."
            ),
            (
                "For every biological unit, the Master received agent states, uncertainties, disagreement, evidence "
                "identifiers and the set of considered future states. It produced an integrated state, Master uncertainty, "
                "a self-observation of expected error, a meta-uncertainty term describing uncertainty in that "
                "self-observation, and an act/defer gate. 'Quantum-inspired' possibility amplitudes were a computational "
                "encoding of normalized future weights; they carried no physical-quantum or consciousness claim."
            ),
        ],
    ),
    (
        "2.4 Recursive scientific observability",
        [
            (
                "In control theory, observability asks whether internal system states can be reconstructed from available "
                "outputs [21]. We use the term more narrowly as an empirical scientific requirement for distributed AI: "
                "the declared agent, Master and recursive self-observer states must be recoverable from logged evidence, "
                "and any claimed coupling to cancer must survive biological-unit null tests. The empirical rank reported "
                "here is not a nonlinear observability Gramian and does not establish controllability."
            ),
            (
                "The complete architecture is termed Global Observability of Distributed Reinforcement States "
                "(GOD-Observability). The acronym denotes coverage across distributed observer states only. It is not a "
                "claim of consciousness, sentience, divinity or physical quantum behaviour. Quantum-inspired amplitudes "
                "are normalized computational weights over considered futures and are tested only for valid encoding and "
                "traceability."
            ),
            (
                "Five time-aligned instability streams were never collapsed into one score: observed cancer future-state "
                "instability; individual-agent instability; inter-agent disagreement; Master-state instability; and "
                "uncertainty in the Master's self-observation. Every O0-O6 event had to satisfy a fail-closed 20-field "
                "contract containing episode and time identifiers, cancer and agent identities, before/after agent and "
                "Master states, action, reward, agent/Master/meta uncertainty, evidence, parent trace, considered futures, "
                "possibility amplitudes, gate status and before/after hypotheses. Missing provenance, empty or non-finite "
                "states, dimension mismatch, invalid uncertainty or inconsistent future weights caused failure."
            ),
            (
                "Three profiles were reported without a post-hoc weighted total. Master Scientific Observability (MSOI) "
                "measured trace completeness, state reconstructability, agent attribution, uncertainty and transition "
                "visibility, and evidence completeness. Recursive Master Observability (RMOI) added agent-influence "
                "identifiability, meta-uncertainty observability and recursive consistency. Scientific Agent Observability "
                "(SAOI) added hypothesis-transition traceability and possibility-state observability. State "
                "reconstructability quantified recovery of the declared Master state from its recorded observables. "
                "Recursive consistency measured agreement between uncertainty/self-observation and realized instability."
            ),
            (
                "Observer-System Instability Coupling (OSIC) was the association between observed biological instability "
                "and Master instability. Observer-Observer/System Coupling (OOSC) preserved pairwise relations among "
                "cancer, agent, disagreement, Master and meta channels. These couplings are diagnostics, not causal "
                "effects. Empirical channel rank quantified whether all five declared streams remained visible in the "
                "current representation; full rank did not imply statistical independence or nonlinear observability."
            ),
        ],
    ),
    (
        "2.5 Offline reinforcement response to instability",
        [
            (
                "A detached reinforcement extension tested how the AI could react after instability became observable. "
                "Using 598 observed calibration transitions from 19 perturbations, it learned reward-conditioned "
                "per-transition weights for ridge, random-forest and nearest-neighbour observers. The policy was a soft "
                "reward-weighted mixture with temperature 0.025. A separate self-error model estimated the expected "
                "Master error and supported an uncertainty-ranked act/defer gate. The 712 transitions from 24 unseen "
                "perturbations remained frozen. Comparators were the original frozen Master, equal ensemble, hard "
                "reward policy, best fixed observer and every individual observer. Reward was a monotone transformation "
                "of normalized forecast error; it was not a clinical reward."
            ),
        ],
    ),
    (
        "2.6 Evaluation of unknown and previously unmodelled features",
        [
            (
                "The declared five-channel hierarchy can expose known unknowns - uncertainty attached to a recorded "
                "state - but it cannot automatically identify unknown unknowns that were absent from the representation. "
                "We therefore defined an unknown-feature candidate as reproducible residual structure that remained after "
                "conditioning observed cancer change on agent instability, disagreement, Master instability and "
                "meta-uncertainty. A candidate would require stability across biological-unit resampling, enrichment in "
                "frozen perturbations or held-out projects, incremental out-of-sample value beyond the declared channels, "
                "and replication in an independent dataset or assay."
            ),
            (
                "Candidate unknown features were not assigned a biological label from model weights alone. They were "
                "classified as representation gaps, observer gaps or possible biological novelty according to whether "
                "they disappeared after adding measurement context, adding an independent observer or obtaining orthogonal "
                "experimental evidence. This protocol prevents unexplained latent structure from being renamed as discovery "
                "without falsification. In the current study it is an evaluation framework; prospective latent-feature "
                "discovery and biological validation remain future experiments."
            ),
        ],
    ),
    (
        "2.7 Statistical analysis",
        [
            (
                "Inference respected the biological unit. Cluster bootstraps and permutations used 2,000 replicates at "
                "the episode, patient or held-out perturbation level; three agent events from one unit were never treated "
                "as independent observations. LINCS nonlinear coupling was tested in a frozen nine-pair specification "
                "using Pearson correlation, Spearman correlation and quantile-normalized mutual information with 1,000 "
                "observed-value permutations. Benjamini-Hochberg adjustment controlled false discovery [12]. TCGA AUROC, "
                "Brier score, calibration and OSIC used 1,000 observed-label permutations. The reinforcement analysis "
                "used held-out perturbation-cluster bootstrap confidence intervals and cluster sign-flip tests. Because "
                "the study is exploratory and multi-source, effect sizes, uncertainty intervals and negative results were "
                "interpreted jointly rather than by a binary significance rule."
            ),
        ],
    ),
    (
        "3. Results",
        [],
    ),
    (
        "3.1 A trace-complete pan-cancer observability system",
        [
            (
                "The consolidated audit evaluated 4,431 real-data transitions: 2,169 DepMap-PRISM events, 2,136 LINCS "
                "events and 126 TCGA outcome events. Strict transition validity, trace completeness, attribution coverage, "
                "uncertainty visibility, evidence provenance, hypothesis-transition traceability and possibility-state "
                "visibility were 1.0 in all systems. The result establishes that the internal decision process was "
                "reconstructable and testable; it does not imply that each forecast was correct. Empirical Master-state "
                "rank was 8/8 dimensions for DepMap-PRISM, 16/16 for LINCS and 1/1 for TCGA."
            ),
            (
                "The three systems showed sharply different scientific profiles. DepMap-PRISM state reconstructability "
                "was 0.999 and recursive consistency 0.956, but cancer-Master OSIC was only 0.271. LINCS reconstructability "
                "fell to 0.223 and recursive consistency to 0.568 while OSIC rose to 0.895. TCGA showed complete state "
                "reconstructability but recursive consistency of 0.436 and OSIC of 0.896. Thus telemetry completeness did "
                "not guarantee predictive stability, and high cancer-AI coupling could coexist with poor reconstruction "
                "or poor self-observation. This separation is the intended falsification mechanism."
            ),
        ],
    ),
    (
        "3.2 Instability was structured, reproducible and significant beyond null worlds",
        [
            (
                "All five instability channels retained empirical rank 5/5 in each architecture. In DepMap-PRISM, "
                "cancer instability correlated with inter-agent disagreement (r=0.252) and Master instability (r=0.271), "
                "whereas its relation to mean individual-agent instability was near zero (r=-0.031). In LINCS, observed "
                "transcriptional change correlated strongly with individual-agent error (r=0.902), disagreement "
                "(r=0.670), Master error (r=0.895) and Master self-observation uncertainty (r=0.611). The channels were "
                "visible but highly collinear (condition number 174.6), preventing an independence claim. In TCGA, "
                "outcome error correlated with mean agent error (r=0.776) and Master error (r=0.896), but self-observation "
                "uncertainty was inversely related to realized error (r=-0.444)."
            ),
            (
                "Cluster-bootstrap 95% intervals for OSIC were 0.203-0.337 for DepMap-PRISM, 0.870-0.915 for LINCS and "
                "0.797-0.949 for TCGA. Across 1,477 independent biological units, all three Master architectures and all "
                "nine observer agents passed structural observability verification. This verifies coverage and "
                "consistency of logged and counterfactual influence. Because no intervention selectively removed an "
                "agent or varied recursion depth, it does not establish causal agent influence or an optimal observer "
                "depth."
            ),
            (
                "The inferential result is stronger than a descriptive correlation. In the frozen LINCS test, "
                "the cancer-change/Master-error relation remained significant across Pearson, Spearman and nonlinear "
                "mutual-information tests after observed-value permutation and false-discovery correction (all empirical "
                "P=0.001; FDR=0.001). The convergence of three dependence measures rejects the narrow null that the "
                "observed coupling is an artifact of marginal value distributions. Combined with positive biological-unit "
                "bootstrap intervals in all three systems, this identifies a reproducible instability field rather than "
                "an arbitrary logging pattern. It does not, by itself, identify a causal biological mechanism."
            ),
        ],
    ),
    (
        "3.3 LINCS exposed how AI reacts to unstable cancer behaviour",
        [
            (
                "The frozen LINCS audit delivered the clearest dynamic test. Large observed changes between matched "
                "time points were accompanied by larger errors in all observers and the Master. Pearson cancer-change/"
                "Master-error coupling was 0.895, Spearman coupling was 0.799 and normalized mutual information was "
                "0.265; all three survived 1,000 observed-value permutations with empirical P=0.001 and FDR=0.001. "
                "This pattern indicates structured forecast failure: the model becomes unstable where the biological "
                "trajectory moves further. It is not proof that the AI has identified a causal mechanism of change."
            ),
            (
                "The self-observer nevertheless contained actionable information. Predicted and actual Master error "
                "correlated at r=0.615 (95% bootstrap interval 0.565-0.663), and high-error detection AUROC was 0.808. "
                "When the system retained only the 20% of transitions with lowest predicted uncertainty, mean absolute "
                "error declined from 0.731 to 0.459. The relevant reaction is therefore not to force a confident answer. "
                "The AI exposes divergence, estimates its likely error and can defer high-risk transitions. Low state "
                "reconstructability of 0.223 remains the central warning: observability made the limitation legible but "
                "did not solve future-state prediction."
            ),
        ],
    ),
    (
        "3.4 Patient-outcome observability revealed a self-monitoring failure",
        [
            (
                "In 42 patients from seven cancer projects absent from fitting and calibration, the two-year "
                "overall-survival Master achieved AUROC 0.677, average precision 0.607 and Brier score 0.215 versus 0.239 "
                "for the fit-prevalence baseline. Observed-label permutation P values were 0.027 for AUROC and 0.036 for "
                "Brier score. These results support the fixed small project holdout but do not constitute external "
                "clinical validation. The sample contained 15 events and 27 non-events, and the project-level split "
                "creates substantial interval uncertainty."
            ),
            (
                "More importantly, the self-error model failed to generalize. Predicted-versus-actual error correlation "
                "was -0.267 (95% interval -0.580 to 0.101), high-error detection AUROC was 0.327 and selective retention "
                "did not monotonically reduce risk. Recursive consistency was only 0.436. A conventional manuscript might "
                "foreground AUROC and leave this failure peripheral. Under observability-first evaluation, the failure is "
                "a primary result: the model can produce a moderately discriminative endpoint while being unable to know "
                "which patient predictions are unsafe."
            ),
            (
                "This discordance is a novel evaluation result rather than merely a limitation. Held-out discrimination "
                "was better than the observed-label null (AUROC P=0.027; Brier P=0.036), yet the self-observer failed its "
                "own task. The same model therefore passed an endpoint test and failed a reliability-of-knowledge test. "
                "Treating these as co-primary endpoints prevents statistically significant prediction from being mistaken "
                "for trustworthy self-monitoring under cancer-project shift."
            ),
        ],
    ),
    (
        "3.5 Cross-dataset pharmacology and rule-collapse validation",
        [
            (
                "Dependency-pharmacology relations were evaluated independently in GDSC. Of 614 metadata-supported "
                "PRISM/CRISPR drug-target pairs, 166 had the same compound in GDSC and a measured CRISPR target. Effect "
                "direction was concordant for 109/147 GDSC1 tests (74.1%), 85/114 GDSC2 tests (74.6%) and 125/166 pooled "
                "tests (75.3%). PRISM-versus-GDSC effect-size correlations were 0.751, 0.754 and 0.775, respectively. "
                "These results support the reproducibility of association structure across independent screens, while "
                "lineage dependence, assay differences and drug promiscuity prevent causal interpretation."
            ),
            (
                "The same observability logic reframed canonical rule collapse. A gene can be strongly required by "
                "CRISPR while a nominally matched drug produces resistance or weak response; expression can imply a "
                "dependency that pharmacology does not realize. Instead of deleting such contradictions as noise, the "
                "pipeline records them as prioritised hypotheses. A rule-collapse state is valuable precisely because "
                "independent observers disagree under known provenance. Experimental perturbation is still required to "
                "separate target biology from off-target pharmacology, lineage context and assay artifacts."
            ),
        ],
    ),
    (
        "3.6 Reinforcement improved observer use but not decisively over the frozen Master",
        [
            (
                "The reward-conditioned mixture assigned mean weights of 0.246 to ridge, 0.339 to random forest and "
                "0.416 to nearest neighbours, selecting nearest neighbours most often. On the frozen holdout, its mean "
                "normalized L2 error was 0.725 and mean reward 0.6111, compared with reward 0.6097 for the frozen Master, "
                "0.6095 for the equal ensemble, 0.6056 for the hard policy and 0.6054 for the best fixed observer. The "
                "learned mixture outperformed the hard policy and individual agents in perturbation-cluster tests."
            ),
            (
                "However, the primary comparison with the frozen Master was small and uncertain: mean reward difference "
                "0.001413, 95% cluster-bootstrap interval -0.000092 to 0.002988 and cluster sign-flip P=0.145. The "
                "reinforcement extension therefore demonstrates a mechanism of reaction - soft observer reweighting, "
                "recursive error estimation and selective deferral - rather than a decisive gain over a well-calibrated "
                "static ensemble. Its 2,136 events were fully trace-complete, self-error correlation was 0.658 and OSIC "
                "was 0.877. The policy remains an offline cancer-cell forecasting experiment, not a treatment policy."
            ),
        ],
    ),
    (
        "3.7 The reinforcement brain exposed a hierarchy of uncertainty, not yet a validated unknown biology",
        [
            (
                "The model observed the proposed hierarchy at the architectural level. Cancer-state change, individual-"
                "observer instability, inter-observer disagreement, Master instability and uncertainty in the Master's "
                "self-observation were present in every trace-complete event and retained full empirical channel rank. "
                "Their unequal correlations, partial correlations, condition numbers, reconstructability and recursive "
                "consistency show that they were not interchangeable labels for one generic confidence variable."
            ),
            (
                "The hierarchy was also functionally expressed. On frozen LINCS perturbations, greater biological change "
                "tracked greater observer and Master error, the self-observer ranked difficult transitions, and selective "
                "retention reduced error. Under TCGA project shift, endpoint discrimination remained above its null while "
                "self-error ranking reversed. The same architecture could therefore distinguish a prediction from its "
                "estimate of reliability and could reveal when that estimate failed. This is evidence for recursive "
                "uncertainty observability, not evidence of consciousness or complete self-knowledge."
            ),
            (
                "What the brain may have learned beyond the declared channels remains unknown. Low LINCS state "
                "reconstructability, strong channel collinearity and cross-system changes in coupling indicate residual "
                "structure that the present observables do not resolve. These residuals are scientifically valuable because "
                "they locate where an unknown feature could exist, but they cannot distinguish missing measurement, domain "
                "shift, model misspecification or novel cancer biology without the evaluation and validation protocol in "
                "Section 2.6. Accordingly, this study demonstrates detection of an unknown-feature boundary rather than "
                "discovery of a specific unknown biological feature."
            ),
        ],
    ),
    (
        "4. Discussion",
        [],
    ),
    (
        "4.1 Addressing the knowledge gap",
        [
            (
                "The central gap in cancer AI is not simply insufficient accuracy. It is the absence of a scientific "
                "instrument panel that shows what the model is doing when cancer behaviour becomes unstable. Ensemble "
                "methods can improve mean performance and calibration [13,14], but averaging can erase the very conflict "
                "that signals alternative biological futures. Post-hoc explanations identify influential features, yet "
                "they do not reconstruct the sequence of agent states, disagreements, gates and self-errors. This study "
                "addresses that gap by making the internal response a prespecified, provenance-linked object of analysis."
            ),
            (
                "The finding that OSIC was high in LINCS and TCGA but much lower in cross-sectional DepMap-PRISM is "
                "biologically and computationally coherent. Dynamic perturbations and patient endpoints contain larger "
                "state changes and unmeasured influences than harmonized cell-line modalities. The Master consequently "
                "fails more where the observed system changes more. Crucially, the direction of interpretation is not "
                "'high coupling equals a good model.' In LINCS, high coupling coexisted with low reconstructability. "
                "OSIC is a sensor of alignment between biological and computational instability; it can reveal an "
                "informative failure surface."
            ),
        ],
    ),
    (
        "4.2 Why this study is different",
        [
            (
                "The framework reverses the usual hierarchy of evaluation. Accuracy is secondary to whether the system "
                "can be inspected, reconstructed, attributed and falsified. This matters because an opaque AUROC can be "
                "seductive. The TCGA experiment illustrates the problem: endpoint discrimination exceeded random-label "
                "nulls, but the self-observer ranked high-risk patients worse than chance. An observability-first system "
                "does not permit the positive metric to cancel the safety failure."
            ),
            (
                "The study also preserves disagreement as data. Cancer biology already teaches that genetic and "
                "non-genetic states can coexist, transition and respond differently to selection [3-6]. A dependency "
                "observer, pharmacology observer and expression observer therefore need not agree. Their disagreement can "
                "mark missing context, assay limitations or a genuine breakdown of a canonical relation. By attaching "
                "evidence identifiers and counterfactual influence to each transition, the framework converts that "
                "ambiguity into a ranked experimental question rather than an averaged residual."
            ),
            (
                "Finally, the study distinguishes reaction from agency. The AI 'reacts' operationally: it changes observer "
                "weights, expands uncertainty, preserves competing hypotheses and may pass a defer gate. This language "
                "does not imply awareness or autonomous clinical intent. The architecture acronym Global Observability "
                "of Distributed Reinforcement States is a systems metaphor only. Its value lies in auditable state "
                "transition, not anthropomorphic interpretation."
            ),
            (
                "Five convergent findings define the novelty. First, a fail-closed event contract made 4,431 distributed "
                "AI transitions completely attributable across three biological settings. Second, five instability channels "
                "remained visible without being collapsed into one confidence score. Third, frozen LINCS perturbations "
                "showed FDR-significant linear, rank and nonlinear coupling between measured biological change and forecast "
                "failure. Fourth, TCGA demonstrated that a statistically supported endpoint can coexist with failed "
                "self-error ranking. Fifth, independent GDSC screens reproduced the direction of 75.3% of pooled "
                "dependency-pharmacology tests with an effect-size correlation of 0.775. No single result establishes "
                "causality or clinical utility; together they establish recursive observability as an empirically testable "
                "evaluation layer rather than a conceptual metaphor."
            ),
        ],
    ),
    (
        "4.3 How the perspective changes",
        [
            (
                "The conceptual breakthrough is a change in the unit of analysis. The conventional unit is a tumour "
                "profile paired with an endpoint; the proposed unit is a coupled trajectory comprising the biological "
                "state, the plurality of accessible futures, the distributed observers attempting to represent them and "
                "the reliability of the Master observing its own response. This shift makes previously discarded model "
                "behaviour—disagreement, policy switching, reconstruction failure and meta-uncertainty—available as "
                "scientific measurements. It does not abolish genes, pathways or response endpoints. It places them inside "
                "a dynamical and recursively observed system in which their meaning can change with history and context."
            ),
            (
                "The operational logic can be stated in the first-person voice of the architecture: 'Cancer has many "
                "possible futures. My agents have many possible interpretations of those futures. I have many possible "
                "ways of integrating those interpretations. I therefore observe not one uncertainty, but a hierarchy of "
                "uncertainties. My task is to distinguish instability in cancer from instability in my observers, "
                "instability in myself, and instability in my observation of myself.' This is a compact systems-level "
                "description, not an attribution of awareness. Each 'my' corresponds to a recorded computational layer, "
                "and each uncertainty must be evaluated independently against observed outcomes, frozen nulls and external "
                "evidence."
            ),
            (
                "The first perspective shift is from classification to future-space. A static model asks which category "
                "best describes the current sample. A future-state model asks which trajectories remain accessible, how "
                "rapidly they branch and which perturbations narrow the malignant set. This aligns computational analysis "
                "with the evolutionary limits of cancer predictability [3,4]. It also changes target prioritization: the "
                "most useful intervention may not be the one with the largest immediate cytotoxic effect, but the one that "
                "collapses escape routes and prevents re-entry into resistant states."
            ),
            (
                "The second shift is from uncertainty as embarrassment to uncertainty as an observable phenotype of the "
                "model-system interaction. If uncertainty tracks observed change beyond null worlds, it carries information "
                "about where the present representation is under strain. That signal can allocate laboratory validation, "
                "trigger additional measurement or require abstention. Selective prediction is valuable when withholding "
                "an answer is safer than issuing an unreliable one [15]. The LINCS risk-coverage result shows the potential; "
                "the TCGA failure shows why the gate itself must be externally validated."
            ),
            (
                "The third shift is epistemic. Negative findings become productive. Low reconstructability identifies the "
                "need for richer temporal, single-cell or spatial measurements; weak self-error calibration identifies a "
                "need for external cohorts and distribution-shift-aware uncertainty; channel collinearity identifies a "
                "need for interventions that separate biological instability from common error magnitude. The system's "
                "purpose is not to appear certain. It is to make the boundary of knowledge explicit enough to test."
            ),
            (
                "This framing introduces a fourth shift: from explaining only known features to auditing the boundary of "
                "the unknown. A stable residual that survives conditioning, resampling, held-out testing and independent "
                "replication becomes a candidate latent feature; a residual that disappears with richer measurement is a "
                "documented representation gap. Both outcomes are informative. The former prioritizes mechanistic "
                "experiments, whereas the latter specifies which measurement or observer was missing. Recursive "
                "observability therefore does not promise to know every learned feature; it makes failure to know a "
                "measurable and falsifiable property of the system."
            ),
            (
                "Taken together, this does not literally make every existing cancer model obsolete. It changes the "
                "standard by which a future-facing cancer AI should be trusted. A model should not be considered complete "
                "because it predicts an endpoint; it should show how alternative observers reached the state, which "
                "evidence supports it, how uncertainty propagates, whether the self-observer is calibrated, and what "
                "happens when the system defers. In that sense, the framework changes the perspective from AI as an answer "
                "generator to AI as an observable scientific instrument."
            ),
            (
                "For experimental cancer science, that instrument creates a bridge between large-scale computation and "
                "falsifiable laboratory design. A region of high observer disagreement can be decomposed into the agents, "
                "features and transitions that created it; the resulting hypotheses can then be tested through target "
                "engagement, lineage-matched perturbation, time-resolved expression or resistant-clone assays. Repeating "
                "the measurement after perturbation asks whether the intervention merely changes the present phenotype or "
                "actually narrows future-state branching. This is a deeper criterion than ranking biomarkers by static "
                "association. It directs scarce experimental resources toward contradictions that survive provenance "
                "checks, cross-dataset replication and observed-value nulls. It also encourages collection of the missing "
                "measurements indicated by a specific failure: longitudinal sampling for poor future reconstruction, "
                "orthogonal pharmacology for dependency-drug collapse, and richer microenvironmental data when cell-line "
                "observers agree but patient behaviour remains unstable."
            ),
            (
                "For AI governance, recursive observability changes what must be documented before a model can influence "
                "care. Reporting only discrimination, calibration and an explanation image is inadequate when the model "
                "cannot reliably estimate its own error under project or population shift. A deployment dossier should "
                "therefore include transition-schema validity, evidence coverage, agent-set completeness, channel "
                "dependence, self-error calibration and risk-coverage behaviour in external cohorts. The act/defer gate "
                "must be treated as a separately validated model with explicit coverage costs, not as a safety label "
                "attached to the predictor. Human review is then triggered by a reproducible state of the system rather "
                "than a vague confidence threshold. This approach cannot remove uncertainty, but it can make uncertainty "
                "auditable, contestable and actionable - the conditions required for responsible translation from "
                "exploratory pan-cancer modelling to prospective decision support."
            ),
        ],
    ),
    (
        "4.4 Limitations and next experiments",
        [
            (
                "Several limitations constrain interpretation. TCGA provided baseline rather than repeated patient-level "
                "molecular measurements, and the held-out outcome set was small. LINCS measured cancer cell lines and 978 "
                "landmark genes, not complete tumours with immune, stromal and spatial context. DepMap, PRISM and GDSC are "
                "large but remain in-vitro resources subject to lineage imbalance, dose and assay differences, target "
                "annotation error and drug polypharmacology. The three architectures were heterogeneous tasks; cross-system "
                "metric contrasts are descriptive rather than randomised comparisons."
            ),
            (
                "Full empirical channel rank does not prove causal separability, and high LINCS collinearity demonstrates "
                "that error-related channels share variance. Logged counterfactual influence is not equivalent to an "
                "interventional ablation. Observer depth was fixed, so recursion-depth benefit and saturation remain "
                "untested. Thresholds were not clinically preregistered, and the reinforcement reward was computational. "
                "No treatment recommendation should be derived from these outputs."
            ),
            (
                "The decisive next study should combine repeated tumour sampling or liquid biopsy with perturbation-aware "
                "single-cell and spatial measurements in an external cohort. Acceptance thresholds for calibration, "
                "selective risk and instability coupling should be preregistered. Agent ablations and controlled "
                "recursion-depth interventions should test causal influence. Prioritized rule-collapse candidates should "
                "undergo orthogonal target-engagement and resistance assays. Only after those steps should a prospective "
                "clinical study test whether an act/defer policy improves decisions."
            ),
        ],
    ),
    (
        "5. Conclusion",
        [
            (
                "This study establishes and statistically evaluates a real-data, pan-cancer implementation of recursive "
                "scientific observability. "
                "Across 4,431 transitions, the architecture made agent states, disagreement, Master dynamics, uncertainty, "
                "self-observation, hypotheses and evidence reconstructable. Dynamic LINCS and TCGA experiments showed that "
                "AI instability can align strongly with unstable or difficult cancer behaviour, while simultaneously "
                "exposing low forecast reconstructability and failed self-error generalization. The reinforcement extension "
                "demonstrated how an observable AI can react through soft observer reweighting and selective deferral, but "
                "did not establish superiority over the frozen Master or clinical benefit. The central significant result "
                "is instead that model-system instability forms a measurable, null-tested failure surface whose "
                "reconstructability and self-monitoring can be independently accepted or rejected."
            ),
            (
                "The resulting perspective is consequential: cancer should be represented not only by its present state "
                "but by its remaining future-space, and AI should be judged not only by the answer it produces but by how "
                "well its internal response can be observed when that future-space becomes unstable. Structured AI "
                "instability is therefore neither automatically biological insight nor disposable noise. It is a testable "
                "signal that can identify uncertainty, prioritize falsification and define when the model must defer. "
                "External longitudinal and interventional validation is now required to determine whether this shift can "
                "ultimately improve cancer control."
            ),
            (
                "The immediate contribution is therefore twofold: the reinforcement brain can observe a hierarchy of "
                "uncertainties already represented in its traces, and it can mark the boundary beyond which its learned "
                "features remain unknown to investigators. Crossing that boundary requires residual testing, independent "
                "replication and biological perturbation. Unknown does not mean unknowable; it means not yet identified by "
                "evidence strong enough to separate new biology from missing context or model error."
            ),
            (
                "The immediate contribution is a measurable standard for that validation. Future studies can now ask "
                "whether an instability signal replicates across institutions, precedes a resistant transition, changes "
                "after a mechanism-specific intervention and improves outcomes when used to request more evidence or "
                "withhold action. Each question has an observable channel, biological unit and null hypothesis. This "
                "turns a broad idea - that cancer's uncertain futures may be reflected in an AI's internal instability - "
                "into a sequence of experiments that can succeed, fail or be refined. That falsifiability, rather than "
                "rhetorical certainty, is how the proposed perspective can produce durable change. It also makes "
                "uncertainty a shared object for transparent multidisciplinary review."
            ),
        ],
    ),
]


REFERENCES = [
    "Hoadley KA, Yau C, Hinoue T, et al. Cell-of-origin patterns dominate the molecular classification of 10,000 tumors from 33 types of cancer. Cell. 2018;173:291-304.e6. doi:10.1016/j.cell.2018.03.022.",
    "Hanahan D. Hallmarks of cancer: new dimensions. Cancer Discovery. 2022;12:31-46. doi:10.1158/2159-8290.CD-21-1059.",
    "McGranahan N, Swanton C. Clonal heterogeneity and tumor evolution: past, present, and the future. Cell. 2017;168:613-628. doi:10.1016/j.cell.2017.01.018.",
    "Greaves M, Maley CC. Clonal evolution in cancer. Nature. 2012;481:306-313. doi:10.1038/nature10762.",
    "Dagogo-Jack I, Shaw AT. Tumour heterogeneity and resistance to cancer therapies. Nature Reviews Clinical Oncology. 2018;15:81-94. doi:10.1038/nrclinonc.2017.166.",
    "Sharma SV, Lee DY, Li B, et al. A chromatin-mediated reversible drug-tolerant state in cancer cell subpopulations. Cell. 2010;141:69-80. doi:10.1016/j.cell.2010.02.027.",
    "Liu J, Lichtenberg T, Hoadley KA, et al. An integrated TCGA pan-cancer clinical data resource to drive high-quality survival outcome analytics. Cell. 2018;173:400-416.e11. doi:10.1016/j.cell.2018.02.052.",
    "Tsherniak A, Vazquez F, Montgomery PG, et al. Defining a cancer dependency map. Cell. 2017;170:564-576.e16. doi:10.1016/j.cell.2017.06.010.",
    "Corsello SM, Nagari RT, Spangler RD, et al. Discovering the anticancer potential of non-oncology drugs by systematic viability profiling. Nature Cancer. 2020;1:235-248. doi:10.1038/s43018-019-0018-6.",
    "Iorio F, Knijnenburg TA, Vis DJ, et al. A landscape of pharmacogenomic interactions in cancer. Cell. 2016;166:740-754. doi:10.1016/j.cell.2016.06.017.",
    "Subramanian A, Narayan R, Corsello SM, et al. A next generation Connectivity Map: L1000 platform and the first 1,000,000 profiles. Cell. 2017;171:1437-1452.e17. doi:10.1016/j.cell.2017.10.049.",
    "Benjamini Y, Hochberg Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. Journal of the Royal Statistical Society Series B. 1995;57:289-300. doi:10.1111/j.2517-6161.1995.tb02031.x.",
    "Lakshminarayanan B, Pritzel A, Blundell C. Simple and scalable predictive uncertainty estimation using deep ensembles. Advances in Neural Information Processing Systems. 2017;30.",
    "Guo C, Pleiss G, Sun Y, Weinberger KQ. On calibration of modern neural networks. Proceedings of the 34th International Conference on Machine Learning. 2017;70:1321-1330.",
    "Kompa B, Snoek J, Beam AL. Second opinion needed: communicating uncertainty in medical machine learning. npj Digital Medicine. 2021;4:4. doi:10.1038/s41746-020-00367-3.",
    "Kendall A, Gal Y. What uncertainties do we need in Bayesian deep learning for computer vision? Advances in Neural Information Processing Systems. 2017;30.",
    "Marusyk A, Almendro V, Polyak K. Intra-tumour heterogeneity: a looking glass for cancer? Nature Reviews Cancer. 2012;12:323-334. doi:10.1038/nrc3261.",
    "Swanton C. Intratumor heterogeneity: evolution through space and time. Cancer Research. 2012;72:4875-4882. doi:10.1158/0008-5472.CAN-12-2217.",
    "Behan FM, Iorio F, Picco G, et al. Prioritization of cancer therapeutic targets using CRISPR-Cas9 screens. Nature. 2019;568:511-516. doi:10.1038/s41586-019-1103-9.",
    "Keenan AB, Jenkins SL, Jagodnik KM, et al. The Library of Integrated Network-Based Cellular Signatures NIH Program: system-level cataloging of human cells response to perturbations. Cell Systems. 2018;6:13-24. doi:10.1016/j.cels.2017.11.001.",
    "Kalman RE. Contributions to the theory of optimal control. Boletin de la Sociedad Matematica Mexicana. 1960;5:102-119.",
]


SUPP_METHODS = [
    (
        "S1. Expanded rationale and scope",
        [
            "The supplement preserves the complete reporting layer behind the main manuscript. The project was restricted to pan-cancer evidence and did not combine oral-resistome, periodontology-only or unrelated disease datasets. The primary endpoint was scientific observability: whether distributed agent states and their relationship to biological instability could be inspected, reconstructed and falsified. Predictive metrics were deliberately placed second.",
            "The biological interpretation is future-state oriented. A cancer state was treated as a present observation embedded in a set of possible transitions. Multi-resolution quantisation produced fine, medium and coarse state tokens, but tokenization did not create biological evidence. The central test remained whether observed state changes aligned with independently recorded agent, disagreement, Master and meta-instability channels beyond rearranged-value null worlds.",
        ],
    ),
    (
        "S2. Provenance and preprocessing",
        [
            "TCGA RNA-seq files were indexed through the GDC public API and balanced at ten open-access STAR-count files per project. Counts and TPM values were harmonised by gene identifier. TCGA-CDR records were linked through patient barcodes, and sample-type labels were retained. The fixed two-year endpoint excluded patients censored before day 730 without a recorded event and handled exact-horizon censoring explicitly.",
            "DepMap expression and CRISPR dependency and PRISM 24Q2 response matrices were harmonised through model identifiers. Drug-target metadata support was audited before dependency-pharmacology comparison. GDSC1 and GDSC2 fitted-response files were mapped with exact Sanger model identifiers and evaluated separately and pooled.",
            "LINCS GSE70138 Level-5 data were checksum verified. Cancer signatures were matched across time using cell, perturbation, dose and perturbation-type keys. Whole perturbation identifiers were withheld before evaluation. This partition prevents signatures from the same perturbation appearing in calibration and test sets even when cell or dose differs.",
        ],
    ),
    (
        "S3. Transition contract and observability hierarchy",
        [
            "The O0-O6 hierarchy records the observed biological state, each agent state, agent-to-agent divergence, the integrated Master state, self-observation and meta-observation. The 20-field contract is listed in Supplementary Table S2. Validation fails closed on missing evidence, non-finite arrays, state-dimension mismatch, negative uncertainty or invalid future probability/amplitude encoding.",
            "MSOI, RMOI and SAOI were kept as profiles because arbitrary aggregation would allow a perfect telemetry field to hide a failed reconstruction or self-observation field. OSIC and OOSC were computed on separated channels. The empirical observability rank ratio reports visible linear dimensions only; it is not presented as a nonlinear control-theory observability Gramian.",
        ],
    ),
    (
        "S4. Statistical safeguards",
        [
            "Resampling units were matched to the source of independence. DepMap episodes, LINCS holdout transitions/perturbations and TCGA patients/projects were not decomposed into independent agent events. Confidence intervals used biological-unit cluster bootstrap. Permutation tests rearranged observed values or labels and Benjamini-Hochberg correction was applied within prespecified families.",
            "The non-ablation verification tests structural coverage, rank, calibration, dependence and selective risk for all three Master architectures and all nine agents. It does not infer causal influence because agents were not experimentally removed and recursion depth was not randomised. The reinforcement comparison used 2,000 perturbation-cluster bootstrap and sign-flip replicates.",
        ],
    ),
    (
        "S5. Reproducibility and claim boundary",
        [
            "All tables below are rendered from the retrieved CSV outputs; all figures are embedded from the checksum-verified PNG package. The local retrieval exactly matched 176 result paths and two model paths in the remote volumes. Model artifacts include the LINCS future-state estimator and reward-conditioned observer policy.",
            "The evidence supports executable recursive observability, structured cancer-AI instability coupling, external pharmacogenomic association replication and an offline observer-reweighting mechanism. It does not establish clinical utility, treatment efficacy, causal biological control, consciousness, physical quantum effects or prospective patient benefit.",
        ],
    ),
]


PRIMARY_TABLE_CAPTIONS = {
    1: "Real-data system inventory and primary coupling.",
    2: "Required fail-closed transition schema.",
    3: "Master Scientific Observability Index profile.",
    4: "Recursive Master Observability Index profile.",
    5: "Scientific Agent Observability Index profile.",
    6: "Separated instability channels by system.",
    7: "Observer-System and observer-observer coupling profile.",
    8: "Agentic divergence summary.",
    9: "Observer-depth profile.",
    10: "Gate influence summary.",
    11: "Secondary TCGA held-out endpoint performance.",
    12: "TCGA project endpoint profile.",
    13: "Rank, state dynamics and multivariate coupling.",
    14: "Non-ablation Master architecture verification.",
    15: "All-agent observability verification.",
    16: "Agent and Master component performance.",
    17: "Clustered observability confidence intervals.",
    18: "Channel separability, dependence and partial correlation.",
    19: "Recursive uncertainty calibration.",
    20: "Selective risk-coverage analysis.",
}

PRIMARY_TABLE_INTERPRETATIONS = {
    1: "The three systems span cross-sectional cell-model biology, measured perturbational change and held-out patient outcomes. The contrast in OSIC across these settings shows that coupling is context-dependent rather than a universal property manufactured by one architecture.",
    2: "The transition schema converts an AI conclusion into an auditable scientific event. Its importance is methodological: a claim without the state transition, uncertainty, provenance and hypothesis change that produced it is treated as an observability failure.",
    3: "Complete trace, attribution and provenance coexist with low LINCS reconstructability. Observability therefore cannot be reduced to logging completeness; it must reveal when the represented Master state is not recoverable from its declared evidence.",
    4: "Recursive consistency separates a system that records self-uncertainty from one whose self-uncertainty is scientifically reliable. The low TCGA value is a substantive negative result that would disappear inside an arbitrary composite score.",
    5: "Scientific-agent observability confirms that hypotheses, evidence and possibility weights are inspectable for every system. These are necessary conditions for falsification, not proof that the hypotheses or quantum-inspired weights are biologically correct.",
    6: "The five channels occupy different ranges and distributions, demonstrating why a single generic uncertainty score would destroy information. The distinct patterns define separate failure surfaces for molecular, perturbational and clinical applications.",
    7: "High cancer–Master coupling in LINCS and TCGA is accompanied by very different cancer–meta relations. The Master can track difficult biology while its self-observer remains well calibrated, weak or directionally wrong; coupling and self-knowledge are separate scientific properties.",
    8: "Pairwise divergence is structured by observer type: modality observers diverge most in DepMap–PRISM, whereas model-family differences dominate LINCS and TCGA. This makes disagreement a localizable experimental signal rather than an undifferentiated ensemble variance.",
    9: "All O0–O6 levels are present, but the strength of state and meta reconstruction varies. The table establishes achieved observer depth while leaving the optimal depth and the value of deeper recursion as open interventional questions.",
    10: "Removing or gating different evidence streams changes the Master by different magnitudes. The result identifies where influence is recorded and testable, but it remains observational because no randomized agent ablation was performed.",
    11: "The held-out TCGA endpoint is moderately discriminative and improves Brier score over prevalence, yet self-error correlation is negative. The clinically important inference is that an apparently useful predictor may still be unable to recognize which patient-level outputs are unsafe.",
    12: "Whole-project partitioning creates genuine biological shift and reveals sparse event counts in several cancers. This strengthens leakage control while explaining the broad uncertainty and the need for external, larger project-level validation.",
    13: "Full empirical rank shows that the declared Master dimensions were visible, not that the underlying nonlinear cancer system was fully observable. Different switching, persistence and total-correlation patterns indicate architecture-specific state dynamics.",
    14: "Every Master architecture satisfies the structural contract and retains five channels, but LINCS has a condition number of 174.6. Full rank therefore coexists with severe collinearity, directly limiting any claim that the channels are causally independent.",
    15: "All nine observers are traceable over complete biological-unit coverage. Unequal uncertainty–transition relations show that structural observability is an architecture property distinct from predictive calibration or biological validity.",
    16: "Component-level performance shows where integration helps and where it merely combines similar errors. The Master improves LINCS error modestly and TCGA performance only within a small internal holdout; these estimates are baselines for causal ablation, not causal attributions.",
    17: "Clustered intervals resample episodes, transitions or patients rather than duplicated agent events. The consistently positive OSIC intervals support structured coupling, while the broad TCGA intervals and low recursive-consistency interval preserve the limits imposed by 42 patients.",
    18: "Partial correlations reveal that raw coupling can be dominated by shared error magnitude, particularly in LINCS. The appropriate novelty claim is a jointly observable instability field, not five independent biological causes.",
    19: "Self-observation is useful in LINCS (r=0.615; high-error AUROC 0.808), weak in DepMap–PRISM and reversed in TCGA. A self-observer must therefore be validated as its own model rather than assumed reliable because the primary predictor performs well.",
    20: "Uncertainty-based retention lowers LINCS error sharply but does not yield a monotonic TCGA curve. Selective deferral is therefore demonstrated as a feasible reaction in perturbational data and simultaneously falsified as a ready clinical safeguard.",
}

REINF_TABLES = [
    ("calibration_policy_tuning.csv", "Reinforcement policy temperature tuning on calibration perturbations."),
    ("holdout_method_evaluation.csv", "Frozen holdout evaluation of the reinforcement Master and comparators."),
    ("cluster_bootstrap_comparisons.csv", "Perturbation-cluster inference for reward differences."),
    ("learned_policy_weights.csv", "Learned observer weights and selection rates."),
    ("selective_risk_coverage.csv", "Reinforcement Master selective risk-coverage profile."),
]

REINF_TABLE_INTERPRETATIONS = [
    "A low-temperature soft mixture had the best cross-fitted calibration reward, supporting graded observer reweighting instead of brittle winner-take-all selection.",
    "The reinforcement Master has the best mean holdout reward and error among reported methods, but the relevant comparison is inferential rather than rank based because the frozen Master is extremely close.",
    "Clustered inference establishes gains over the hard policy and individual observers but not over the frozen Master or equal ensemble. The novelty is an observable reaction mechanism, not a decisive performance breakthrough.",
    "Nearest neighbours receive the largest mean weight and selection rate, but all observers retain non-zero influence. The Master reacts to local transition context rather than permanently declaring one model universally superior.",
    "Realized error falls as high predicted self-error transitions are deferred. This validates a selective-risk mechanism on unseen perturbations, while its translational value remains untested outside cancer-cell perturbation data.",
]

ALL_FIGURES = [
    (WORKFLOW_FIG, "Recursive observability workflow for unstable cancer futures.", "The workflow changes the unit of analysis from a tumour-prediction pair to a coupled trajectory containing the biological state, the possible futures, the distributed observers and the reliability of the Master observing itself. This is the conceptual bridge between cancer future-space and executable recursive observability."),
    (PRIMARY_FIG / "figure_01_system_inventory.png", "Real-data system inventory.", "Three complementary biological scales contribute 4,431 trace-complete transitions. The figure establishes the evidence topology on which all later claims depend and prevents pan-cancer breadth from being confused with repeated measurements of one system."),
    (PRIMARY_FIG / "figure_02_msoi_heatmap.png", "MSOI profile across the three real-data systems.", "Perfect trace and provenance coverage do not erase the low LINCS state reconstruction. The profile turns this failure into a visible dimension rather than allowing complete telemetry to masquerade as complete understanding."),
    (PRIMARY_FIG / "figure_03_rmoi_heatmap.png", "RMOI profile across the three real-data systems.", "Recursive consistency varies sharply despite universal visibility of uncertainty. The figure demonstrates that observing a self-estimate and trusting that self-estimate are different requirements."),
    (PRIMARY_FIG / "figure_04_saoi_heatmap.png", "SAOI profile across the three real-data systems.", "Evidence, hypotheses, influence and possibility amplitudes remain traceable across architectures. This creates a falsifiable scientific record without implying that the recorded hypotheses are true."),
    (PRIMARY_FIG / "figure_05_instability_channel_distributions.png", "Five non-collapsed instability-channel distributions.", "The channels have different scales and shapes across systems. Their separation is the core methodological novelty because it prevents biological instability, observer error, disagreement and meta-error from being silently averaged into one confidence number."),
    (PRIMARY_FIG / "figure_06_oosc_heatmap.png", "Pairwise observer–observer/system coupling matrix.", "LINCS and TCGA show strong cancer–Master alignment but opposite or weaker meta-level relations. Thus a Master may recognize that a case is difficult while still failing to know whether its own error estimate is reliable."),
    (PRIMARY_FIG / "figure_07_agentic_divergence_matrices.png", "Agentic divergence matrices.", "Disagreement is structured by modality and model family rather than randomly distributed. This supports using divergence to localize a contradiction for experimental follow-up instead of discarding it as ensemble noise."),
    (PRIMARY_FIG / "figure_08_observer_depth_profiles.png", "Observed profiles across the O0–O6 hierarchy.", "The full recursive hierarchy is instrumented, but reconstruction weakens at different depths in LINCS and TCGA. The result defines an achieved observer depth, not evidence that infinite or deeper recursion is beneficial."),
    (PRIMARY_FIG / "figure_09_gate_influence_distributions.png", "Counterfactual gate and observer-influence distributions.", "Evidence streams alter the Master by unequal amounts, making influence inspectable at each transition. Causal influence still requires controlled removal or intervention because these are logged counterfactual diagnostics."),
    (PRIMARY_FIG / "figure_10_tcga_roc_calibration.png", "TCGA held-out discrimination and calibration.", "Moderate discrimination and Brier improvement coexist with failed self-error ranking. The figure is a direct warning that a favourable endpoint curve cannot certify that the model recognizes unsafe patient predictions."),
    (PRIMARY_FIG / "figure_11_lincs_cancer_master_coupling.png", "LINCS observed cancer-change and Master-error coupling.", "Large measured transcriptomic transitions coincide with large Master errors across unseen perturbations. The AI instability is therefore structured by biological change, but its meaning is forecast limitation—not causal mastery of the transition."),
    (PRIMARY_FIG / "figure_12_tcga_project_outcomes.png", "TCGA project-level endpoint composition and predictions.", "Holding out complete cancer projects exposes lineage shift and sparse outcomes that patient-random splits can conceal. The heterogeneity justifies conservative intervals and external validation."),
    (PRIMARY_FIG / "figure_13_master_architecture_verification.png", "Structural verification of all three Master architectures.", "All Masters satisfy fail-closed structural observability and retain five visible channels. The high LINCS condition number preserves the critical distinction between visibility and independence."),
    (PRIMARY_FIG / "figure_14_all_agent_observability_verification.png", "Observability verification for all nine observers.", "Every observer is attributable over complete biological-unit coverage. This establishes a common audit standard across heterogeneous models while leaving their scientific accuracy to separate tests."),
    (PRIMARY_FIG / "figure_15_clustered_observability_intervals.png", "Biological-unit clustered confidence intervals.", "OSIC remains positive when independence is assigned to the episode, transition or patient rather than duplicated agent events. The wide patient intervals make the small TCGA evidence boundary visible."),
    (PRIMARY_FIG / "figure_16_recursive_uncertainty_calibration.png", "Calibration of recursive self-error signals.", "The self-observer is informative for LINCS but reverses direction in TCGA. Recursive self-monitoring is therefore a separately falsifiable model component, not a safety property inherited from the predictor."),
    (PRIMARY_FIG / "figure_17_channel_partial_correlations.png", "Partial correlations among separated instability channels.", "Many raw associations shrink or reverse after conditioning on the remaining channels. The novel object is a coupled instability field; causal independence cannot be inferred from full channel rank."),
    (PRIMARY_FIG / "figure_18_selective_risk_coverage.png", "Selective risk–coverage across systems.", "Deferring high-uncertainty LINCS transitions improves retained accuracy, whereas TCGA risk is not monotonic. The act/defer concept is supported in one domain and falsified as a ready clinical safeguard in another."),
    (PRIMARY_FIG / "figure_19_agent_master_component_performance.png", "Observer and Master component performance.", "Integration provides context-dependent, generally modest gains. The figure prevents the observability advance from being misreported as universal predictive superiority."),
    (REINF_FIG / "figure_01_holdout_reward_comparison.png", "Frozen-holdout reward and error comparison.", "The reinforcement Master ranks first on mean reward, yet its gain over the frozen Master is not statistically decisive. Its contribution is an observable response policy rather than proof of an optimal learner."),
    (REINF_FIG / "figure_02_learned_observer_weights.png", "Learned per-transition observer weights.", "Nearest neighbours are favoured most often while ridge and random forest remain active. Soft weighting preserves plural evidence and avoids the instability of hard model selection."),
    (REINF_FIG / "figure_03_master_self_observation.png", "Predicted versus realized reinforcement-Master error.", "The positive self-error relation shows that the recursively observed Master can rank difficult perturbational transitions sufficiently to support selective action in this offline setting."),
    (REINF_FIG / "figure_04_selective_risk.png", "Selective risk for the reinforcement-Master act/defer gate.", "Error declines as coverage is reduced from all transitions to the lowest-risk fifth. The operational inference is that instability can trigger abstention, not that the policy is clinically safe or therapeutically effective."),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=BLACK)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def set_run_font(run, name="Times New Roman", size=None, bold=None, italic=None, color=BLACK) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document, supplement: bool = False) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 12, BLACK, 14, 6),
        ("Heading 2", 12, BLACK, 10, 4),
        ("Heading 3", 12, BLACK, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor.from_string(BLACK)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.keep_with_next = False

    header = section.header.paragraphs[0]
    header.text = "SUPPLEMENTARY MATERIAL | " + SHORT_TITLE if supplement else SHORT_TITLE
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.runs[0], size=9, bold=False, color=BLACK)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title_block(doc: Document, subtitle: str, word_count: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run(TITLE), size=16, bold=True, color=BLACK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run(subtitle), size=12, italic=True, color=BLACK)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    meta = "Article and Position Paper | Pan-cancer computational study | Blind-review manuscript"
    if word_count:
        meta += " | " + word_count
    set_run_font(p.add_run(meta), size=10, bold=False, color=BLACK)


def add_heading(doc: Document, text: str) -> None:
    if re.match(r"^\d+\.\d+", text):
        level = 2
    else:
        level = 1
    doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    p.add_run(text)


def add_caption(doc: Document, label: str, text: str, keep_with_next: bool = False) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = keep_with_next
    p.paragraph_format.space_before = Pt(4 if keep_with_next else 3)
    p.paragraph_format.space_after = Pt(4 if keep_with_next else 9)
    r = p.add_run(f"{label}. ")
    r.bold = True
    p.add_run(text)


def add_interpretation(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run("Significance and inference. ")
    set_run_font(r, size=11, bold=True, color=BLACK)
    r = p.add_run(text)
    set_run_font(r, size=11, color=BLACK)


def add_figure(doc: Document, path: Path, label: str, caption: str, width_in: float, max_height_in: float = 6.15) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    with Image.open(path) as source_image:
        aspect = source_image.width / source_image.height
    final_width = min(width_in, max_height_in * aspect)
    shape = p.add_run().add_picture(str(path), width=Inches(final_width))
    shape._inline.docPr.set("title", label)
    shape._inline.docPr.set("descr", caption)
    add_caption(doc, label, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int], font_size=8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = str(value)
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=font_size, bold=True, color=BLACK)
    for row_values in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i < 2 else WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run_font(run, size=font_size, color=BLACK)
    set_table_geometry(table, widths_dxa, 120)


def add_main_table_1(doc: Document) -> None:
    headers = ["System", "Biological units", "Events", "SR", "RC", "OSIC"]
    rows = [
        ["DepMap/PRISM Master", "723 models", "2,169", "0.999", "0.956", "0.271"],
        ["LINCS future state", "712 transitions", "2,136", "0.223", "0.568", "0.895"],
        ["TCGA patient outcome", "42 patients", "126", "1.000", "0.436", "0.896"],
    ]
    add_caption(doc, "Table 1", "Primary recursive observability results across three complementary real-data systems. SR, state reconstructability; RC, recursive consistency; OSIC, observer-system instability coupling.", keep_with_next=True)
    add_table(doc, headers, rows, [3100, 1750, 1100, 1100, 1100, 1210], 8.5)
    add_interpretation(doc, "The decisive and novel pattern is dissociation: high coupling can coexist with low reconstruction or failed recursive consistency. Because positive biological-unit intervals recur across all three systems and the LINCS relation survives three FDR-controlled null tests, the result supports a reproducible instability field rather than generic telemetry noise. The framework therefore makes scientific failure measurable instead of allowing complete logging or moderate prediction to conceal it.")


def add_main_table_2(doc: Document) -> None:
    headers = ["Setting", "Instability signal", "AI reaction", "Boundary"]
    rows = [
        ["DepMap/PRISM", "Modest cancer-Master coupling", "Preserve modality disagreement and rule-collapse hypotheses", "Association, not target causality"],
        ["LINCS", "Large biological change tracks forecast error", "Raise self-risk, reweight observers, selectively defer", "Cell-line forecasting; low reconstructability"],
        ["TCGA", "Outcome error couples to Master error", "Expose failed self-error ranking instead of trusting AUROC alone", "Small internal project holdout"],
        ["Offline reinforcement", "Observer reward varies by transition", "Soft per-transition weighting plus act/defer gate", "No decisive gain over frozen Master"],
    ]
    add_caption(doc, "Table 2", "How the observable AI responds to unstable cancer behaviour without converting uncertainty into an unsupported clinical claim.", keep_with_next=True)
    add_table(doc, headers, rows, [1800, 2600, 3000, 1960], 8.2)
    add_interpretation(doc, "Reaction is operationally defined as preserving alternatives, changing observer influence, requesting evidence or deferring. It does not denote awareness, clinical autonomy or control of cancer.")


def add_main_table_3(doc: Document) -> None:
    headers = ["Question", "Conventional frame", "Proposed frame", "Falsification criterion"]
    rows = [
        ["What is cancer?", "A present molecular state or endpoint", "A history-dependent field of accessible futures", "No reproducible transition structure beyond static labels"],
        ["What is model disagreement?", "Noise to average or calibrate", "A separated observer-state signal", "Coupling disappears under frozen nulls or external data"],
        ["What makes the Master trustworthy?", "Mean predictive performance", "Reconstructability plus calibrated self-observation", "Self-error cannot rank unsafe outputs"],
        ["How should AI react?", "Return the most likely answer", "Reweight, falsify, request measurement or defer", "Risk does not improve when high self-risk outputs are withheld"],
        ["What is an unknown learned feature?", "An unexplained latent weight or embedding", "A reproducible residual candidate beyond declared instability channels", "No held-out increment, resampling stability or independent replication"],
    ]
    add_caption(doc, "Table 3", "Position statement and explicit falsifiers for a future-facing cancer AI.", keep_with_next=True)
    add_table(doc, headers, rows, [1900, 2250, 2700, 2510], 8.2)
    add_interpretation(doc, "The proposal changes the standard of evidence without making itself unfalsifiable. Every conceptual claim is paired with a result that could invalidate it, which is the basis for treating recursive observability as a scientific method rather than a metaphor.")


def narrative_word_count() -> int:
    words = []
    for heading, paragraphs in SECTIONS:
        if heading.startswith(("1.", "2.", "3.", "4.", "5.")):
            words.append(heading)
            words.extend(paragraphs)
    return len(re.findall(r"\b[\w'-]+\b", " ".join(words)))


def build_main() -> None:
    doc = Document()
    configure_styles(doc, supplement=False)
    count = narrative_word_count()
    add_title_block(
        doc,
        "Pan-cancer evidence for turning distributed AI instability into a falsifiable biological sensing surface",
        f"Main-text word count: {count}",
    )

    add_heading(doc, "Abstract")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run(ABSTRACT)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("Keywords: "), bold=True, color=BLACK)
    p.add_run("pan-cancer; recursive observability; multi-agent AI; cancer instability; LINCS; TCGA; DepMap; uncertainty; reinforcement learning; selective prediction")

    for heading, paragraphs in SECTIONS:
        add_heading(doc, heading)
        for text in paragraphs:
            add_body(doc, text)
        if heading == "1. Introduction":
            add_figure(
                doc,
                WORKFLOW_FIG,
                "Figure 1",
                "Recursive observability workflow. Real pan-cancer observations enter independent molecular, perturbational and clinical observers. Five synchronized instability streams remain separate while the Master Reinforcement AI integrates observer states under fail-closed O0-O6 telemetry. MSOI, RMOI, SAOI, OSIC, OOSC and biological-unit null tests determine whether internal instability is reconstructable and coupled to cancer change. The AI reacts by preserving alternatives, reweighting observers, requesting evidence or deferring; these are research actions, not clinical treatment decisions.",
                6.35,
            )
        elif heading == "3.1 A trace-complete pan-cancer observability system":
            add_main_table_1(doc)
            add_figure(
                doc,
                PRIMARY_FIG / "figure_01_system_inventory.png",
                "Figure 2",
                "Real-data architecture inventory. Three complementary systems preserve synchronized cancer, agent, disagreement, Master and meta-observation streams. All counts refer to observed public pan-cancer data.",
                6.35,
            )
            add_figure(
                doc,
                PRIMARY_FIG / "figure_13_master_architecture_verification.png",
                "Figure 3",
                "Architecture-level verification. All three Masters satisfy the fail-closed structural contract and retain five visible channels; the LINCS condition number shows that visibility is not causal independence.",
                6.25,
            )
        elif heading == "3.2 Instability was structured, reproducible and significant beyond null worlds":
            add_figure(
                doc,
                PRIMARY_FIG / "figure_05_instability_channel_distributions.png",
                "Figure 4",
                "Separated instability-channel distributions. Distinct channel behaviour prevents a generic uncertainty score from hiding reconstruction, disagreement or self-observation failures.",
                6.35,
            )
            add_figure(
                doc,
                PRIMARY_FIG / "figure_17_channel_partial_correlations.png",
                "Figure 5",
                "Partial correlations among the five channels. Dependence changes after conditioning, supporting a coupled instability field while rejecting an unsupported claim of causal independence.",
                6.25,
            )
        elif heading == "3.3 LINCS exposed how AI reacts to unstable cancer behaviour":
            add_figure(
                doc,
                PRIMARY_FIG / "figure_11_lincs_cancer_master_coupling.png",
                "Figure 6",
                "Measured LINCS cancer-state change versus Master error on unseen perturbations. Strong coupling identifies a structured forecast-failure surface rather than successful causal prediction.",
                6.25,
            )
            add_figure(
                doc,
                PRIMARY_FIG / "figure_16_recursive_uncertainty_calibration.png",
                "Figure 7",
                "Recursive self-error calibration. The LINCS self-observer can identify difficult transitions, whereas the TCGA self-observer reverses direction, making self-knowledge a separately falsifiable component.",
                6.25,
            )
        elif heading == "3.4 Patient-outcome observability revealed a self-monitoring failure":
            add_figure(
                doc,
                PRIMARY_FIG / "figure_10_tcga_roc_calibration.png",
                "Figure 8",
                "TCGA held-out project performance and calibration. Moderate discrimination coexists with failed self-error generalization, demonstrating why performance alone is insufficient.",
                6.2,
            )
        elif heading == "3.6 Reinforcement improved observer use but not decisively over the frozen Master":
            add_figure(
                doc,
                REINF_FIG / "figure_01_holdout_reward_comparison.png",
                "Figure 9",
                "Reward and error comparison on 712 transitions from 24 frozen perturbations. The reward-conditioned mixture exceeds individual observers but is not decisively superior to the frozen Master.",
                6.2,
            )
        elif heading == "4.3 How the perspective changes":
            add_main_table_2(doc)
            add_main_table_3(doc)

    add_heading(doc, "References")
    for idx, ref in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.first_line_indent = Inches(-0.28)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(f"{idx}. {ref}")

    add_heading(doc, "Data and code availability")
    add_body(doc, "All manuscript figures, tables, reports, traces, predictions, checksums and trained model artifacts are preserved in the accompanying retrieved results package. Raw and processed public-data volumes remain in the declared computational workspace. The manuscript reports immutable artifact-level results and the supplement maps every rendered item to its source filename.")
    add_heading(doc, "Ethics statement")
    add_body(doc, "This secondary computational analysis used de-identified public data and did not recruit participants or perform a prospective intervention. Users should apply their institutional requirements to any future clinical or patient-level validation.")
    add_heading(doc, "Competing interests and funding")
    add_body(doc, "No competing-interest or funding information was supplied for this blind-review draft. These declarations should be completed by the submitting authors before journal submission.")

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Pan-cancer recursive scientific observability manuscript"
    doc.core_properties.keywords = "pan-cancer, observability, multi-agent AI, cancer instability"
    doc.core_properties.author = "Blind review"
    doc.save(MAIN_DOCX)


def read_csv_table(path: Path) -> tuple[list[str], list[list[str]]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.reader(handle)
        all_rows = list(reader)
    if not all_rows:
        return ["status"], [["empty"]]
    headers = all_rows[0]
    rows = all_rows[1:]
    return headers, rows


def format_value(value: str) -> str:
    value = str(value)
    try:
        number = float(value)
        if math.isfinite(number):
            if abs(number) >= 1000 and number.is_integer():
                return f"{int(number):,}"
            if abs(number) < 0.001 and number != 0:
                return f"{number:.3g}"
            if abs(number) < 100:
                return f"{number:.4f}".rstrip("0").rstrip(".")
    except ValueError:
        pass
    return value.replace("|", " / ")


def column_widths(headers: list[str], rows: list[list[str]], total_dxa: int) -> list[int]:
    weights = []
    for i, header in enumerate(headers):
        sample = [header] + [row[i] if i < len(row) else "" for row in rows[:40]]
        length = max(7, min(26, max(len(str(x)) for x in sample)))
        weights.append(math.sqrt(length))
    raw = [max(900, int(total_dxa * w / sum(weights))) for w in weights]
    scale = total_dxa / sum(raw)
    widths = [int(x * scale) for x in raw]
    widths[-1] += total_dxa - sum(widths)
    return widths


def add_csv_as_split_tables(doc: Document, path: Path, label: str, caption: str) -> None:
    headers, rows = read_csv_table(path)
    rows = [[format_value(v) for v in row] for row in rows]
    max_cols = 4
    if len(headers) <= max_cols:
        groups = [list(range(len(headers)))]
    else:
        anchors = list(range(min(2, len(headers))))
        remaining = list(range(len(anchors), len(headers)))
        groups = [anchors + remaining[i:i + (max_cols - len(anchors))] for i in range(0, len(remaining), max_cols - len(anchors))]
    for panel, indices in enumerate(groups, start=1):
        panel_label = label if len(groups) == 1 else f"{label}, panel {chr(64 + panel)} of {len(groups)}"
        panel_caption = caption if len(groups) == 1 else caption + " Wide tables are split into repeated identifier panels for readability."
        add_caption(doc, panel_label, panel_caption, keep_with_next=True)
        h = [headers[i] for i in indices]
        r = [[row[i] if i < len(row) else "" for i in indices] for row in rows]
        widths = column_widths(h, r, 9360)
        add_table(doc, h, r, widths, 8.2)
        doc.add_paragraph().paragraph_format.space_after = Pt(3)


def build_supplement() -> None:
    doc = Document()
    configure_styles(doc, supplement=True)
    add_title_block(doc, "Supplementary methods, complete result tables and complete 24-figure evidence atlas")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("Supplementary scope: "), bold=True, color=BLACK)
    p.add_run("This file contains the expanded reproducibility record supporting the main manuscript. All 25 result tables are rendered directly from retrieved CSV artifacts, and the new workflow plus all 23 retrieved result figures are embedded as a complete 24-figure evidence atlas. Every item is followed by a specific significance and inference statement.")

    for heading, paragraphs in SUPP_METHODS:
        add_heading(doc, heading)
        for text in paragraphs:
            add_body(doc, text)

    doc.add_page_break()
    add_heading(doc, "S6. Complete supplementary tables")
    add_body(doc, "Supplementary Tables S1-S20 reproduce the primary publication tables. Tables S21-S25 reproduce the offline reinforcement Master outputs. Values marked not_applicable were intentionally undefined rather than missing.")
    for idx in range(1, 21):
        matches = list(PRIMARY_TAB.glob(f"table_{idx:02d}_*.csv"))
        if len(matches) != 1:
            raise RuntimeError(f"Expected one primary table {idx}, found {matches}")
        add_csv_as_split_tables(doc, matches[0], f"Supplementary Table S{idx}", PRIMARY_TABLE_CAPTIONS[idx] + f" Source: {matches[0].name}.")
        add_interpretation(doc, PRIMARY_TABLE_INTERPRETATIONS[idx])
        if idx in {5, 10, 15, 20}:
            doc.add_page_break()
    for offset, (filename, caption) in enumerate(REINF_TABLES, start=21):
        add_csv_as_split_tables(doc, REINF / filename, f"Supplementary Table S{offset}", caption + f" Source: {filename}.")
        add_interpretation(doc, REINF_TABLE_INTERPRETATIONS[offset - 21])

    doc.add_page_break()
    add_heading(doc, "S7. Complete figure evidence atlas")
    add_body(doc, "The recursive-observability workflow, all nineteen primary observability figures and all four reinforcement-Master figures are reproduced below. Central figures are intentionally repeated from the main manuscript so the supplement remains a self-contained 24-figure evidence atlas.")
    for idx, (path, caption, interpretation) in enumerate(ALL_FIGURES, start=1):
        add_figure(doc, path, f"Supplementary Figure S{idx}", caption + f" Source: {path.name}.", 6.25)
        add_interpretation(doc, interpretation)
        if idx != len(ALL_FIGURES):
            doc.add_page_break()

    add_heading(doc, "S8. Artifact map")
    artifact_rows = [
        ["Complete local retrieval", "retrieved_results_2026-08-25", "176 result files and 2 model files; exact remote/local path match"],
        ["Primary publication package", "results/observability_primary/publication_package", "19 primary PNG/SVG figure pairs; 20 CSV/HTML table pairs"],
        ["Reinforcement package", "results/reinforcement_master_brain", "4 PNG/SVG figure pairs; 5 CSV tables; report, events and predictions"],
        ["Model artifacts", "models", "LINCS future-state and reward-conditioned observer-policy Joblib files"],
        ["Integrity manifest", "SHA256_MANIFEST.csv", "All locally retrieved artifacts verified without failure"],
    ]
    add_table(doc, ["Artifact", "Relative location", "Contents"], artifact_rows, [1800, 3360, 4200], 8.2)

    doc.core_properties.title = TITLE + " - Supplementary Material"
    doc.core_properties.subject = "Complete supplementary methods, tables and figures"
    doc.core_properties.author = "Blind review"
    doc.save(SUPP_DOCX)


def main() -> None:
    OUT.mkdir(exist_ok=True)
    QA.mkdir(exist_ok=True)
    build_main()
    build_supplement()
    print(f"main={MAIN_DOCX}")
    print(f"supplement={SUPP_DOCX}")
    print(f"main_text_words={narrative_word_count()}")


if __name__ == "__main__":
    main()
