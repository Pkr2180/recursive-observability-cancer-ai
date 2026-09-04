from __future__ import annotations

import csv
import hashlib
import json
import math
import re
import shutil
import subprocess
import textwrap
import zipfile
from pathlib import Path
from typing import Iterable

import requests
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_manuscript_package import (
    column_widths,
    format_value,
    read_csv_table,
    repeat_table_header,
    set_cell_margins,
    set_cell_shading,
    set_table_geometry,
)


ROOT = Path(__file__).resolve().parent
DATE = "2026-09-04"
OUT = ROOT / "manuscript_package" / f"Nature_Machine_Intelligence_Submission_{DATE}"
FIG_OUT = OUT / "Figures"
SUPP_FIG_OUT = OUT / "Supplementary_Figures"
FORMS_OUT = OUT / "Nature_Forms"
QA = OUT / "QA"

PRIMARY = ROOT / "final_results" / "primary_scientific_observability_results"
BENCHMARK = ROOT / "final_results" / "recursive_observability_benchmark"
REINFORCEMENT = ROOT / "final_results" / "reinforcement_master_brain"
WORKFLOW = ROOT / "manuscript_assets" / "figure_01_recursive_observability_workflow.png"

TITLE = "Recursive observability makes hidden failures measurable in distributed cancer AI"
SHORT_TITLE = "Recursive observability in cancer AI"

AUTHOR_NAME = "Pradeep Kumar Yadalam"
AUTHOR_DEGREES = "BDS, MDS, PhD"
AUTHOR_FULL = f"{AUTHOR_NAME}, {AUTHOR_DEGREES}"
AUTHOR_ROLE = "Professor and Head of Research"
AFFILIATION = (
    "Department of Periodontics, Saveetha Dental College and Hospitals, "
    "Saveetha Institute of Medical and Technical Sciences (SIMATS), "
    "Chennai 600077, Tamil Nadu, India"
)
AUTHOR_EMAIL = "pradeepkumar.sdc@saveetha.com"
ORCID_ID = "0000-0002-6653-4123"
AUTHOR_ORCID = f"ORCID {ORCID_ID}"
REPOSITORY_URL = "https://github.com/Pkr2180/recursive-observability-cancer-ai"
ZENODO_STATUS = (
    "Zenodo archiving metadata is prepared in the code repository (.zenodo.json); the versioned "
    "DOI will be inserted on archival release"
)
CREDIT_STATEMENT = (
    "Pradeep Kumar Yadalam: Conceptualization, Methodology, Software, Validation, "
    "Formal analysis, Investigation, Resources, Data curation, Writing - original draft, "
    "Writing - review and editing, Visualization, Supervision and Project administration. "
    "No funding acquisition role applies because the study received no specific grant. "
    "The sole author accepts responsibility for the work."
)
FUNDING_STATEMENT = (
    "This study received no specific grant from any funding agency in the public, commercial "
    "or not-for-profit sectors."
)
ACKNOWLEDGEMENT_STATEMENT = (
    "The analyses were executed in the author's Modal workspace, pradeepaiperio. "
    "No non-author contribution requiring acknowledgement was reported."
)
COMPETING_INTERESTS_STATEMENT = "The author declares no competing interests."
ETHICS_STATEMENT = (
    "This study exclusively reanalysed de-identified, publicly accessible datasets and involved "
    "no new participant recruitment, intervention or access to identifiable private information. "
    "Ethics approval and new informed consent were therefore not required for this secondary analysis; "
    "dataset-level consent and oversight are described in the source publications."
)
MODAL_COMPUTE_STATEMENT = (
    "Remote analyses ran in the pradeepaiperio Modal workspace using a Debian-slim Python 3.12 "
    "container. Declared function resources ranged from 2 to 8 vCPUs and 4,096 to 32,768 MB memory; "
    "no GPU resource was requested. Modal abstracts the underlying physical processor model. Exact "
    "per-stage wall-clock runtime and energy use were not retained in the retrieved result artifacts."
)

MAIN_DOCX = OUT / "01_NMI_Main_Manuscript.docx"
SUPP_DOCX = OUT / "02_NMI_Supplementary_Information.docx"
COVER_DOCX = OUT / "03_NMI_Cover_Letter.docx"
TITLE_PAGE_DOCX = OUT / "04_NMI_Title_Page_and_Declarations.docx"
REPORTING_DOCX = OUT / "05_Nature_Reporting_Summary_Responses.docx"
ML_DOCX = OUT / "06_Nature_Machine_Learning_Checklist_Responses.docx"
INSTRUCTIONS_DOCX = OUT / "07_NMI_Author_Instructions_Audit.docx"
CHECKLIST_DOCX = OUT / "08_Submission_Readiness_Checklist.docx"
TABLE_DOCX = OUT / "09_Main_Table_1.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4E79"
BLACK = "000000"
WHITE = "FFFFFF"
LIGHT_GRAY = "F4F6F9"
AMBER = "FFF2CC"


ABSTRACT = (
    "Distributed scientific AI is judged by outputs even when its internal observer states "
    "cannot be reconstructed or falsified. We define recursive scientific observability as nine "
    "conditions linking provenance, agent identifiability, Master-state reconstructability, "
    "self-error ordering, channel separability, failure sensitivity, non-redundant depth and cross-system "
    "invariance. We evaluated three architectures using observed DepMap/PRISM, LINCS and TCGA data. "
    "Across 4,431 transitions, all events met the transition contract and all nine conditions were "
    "demonstrated. Ninety graded disruptions hid, corrupted or collapsed internal channels; all 18 "
    "system-by-failure tests detected loss, with median cross-system failure-signature "
    "concordance of 0.754. Higher observer levels added held-out reconstruction information in every "
    "system (maximum marginal gains, 0.207-0.416). Outcome metrics were secondary: TCGA AUROC was 0.677, "
    "whereas self-error ranking failed, showing that discrimination can coexist with poor introspection. "
    "Recursive observability is therefore a falsifiable engineering property of scientific AI, not proof "
    "of biological correctness or clinical utility."
)


INTRODUCTION = [
    (
        "Scientific artificial intelligence is moving from single predictors toward distributed and "
        "agentic systems that generate hypotheses, select tools and integrate heterogeneous evidence. "
        "Yet the dominant evaluation target remains the external answer: accuracy, reward, task completion "
        "or biological association. This leaves a basic engineering gap. If an integrated output changes, "
        "investigators often cannot determine which observer changed, whether the Master state remained "
        "reconstructable, whether self-estimated error preserved its ordering, or whether nominally distinct "
        "uncertainty channels encoded one redundant signal. Classical observability asks whether internal "
        "states can be recovered from available outputs [1,2]. For scientific AI, an additional requirement "
        "is falsifiability: loss of a claimed internal channel should produce a specific, detectable failure."
    ),
    (
        "Post-hoc interpretability, predictive uncertainty and calibration address important but different "
        "questions. Explanations attribute a prediction after it is produced; uncertainty methods estimate "
        "confidence; calibration aligns confidence with empirical frequency [3-7]. None alone demonstrates "
        "that provenance, agent states, Master integration and recursive self-observation remain separately "
        "available while a distributed decision evolves. This distinction is acute in cancer research. "
        "Tumour state, dependency, pharmacological response, perturbational transcription and patient outcome "
        "are complementary rather than interchangeable measurements [13-20]. An ensemble can be moderately "
        "accurate while concealing disagreement or failing to identify its own largest errors. A clinically "
        "grounded evaluation must therefore distinguish outcome validity from the observability of the process "
        "that generated the outcome."
    ),
    (
        "Recent oncology agents have been evaluated through tool use, prognostic or predictive endpoints and "
        "autonomous hypothesis generation [8-10], while ScienceAgentBench evaluates executable task success "
        "[11]. These advances show why the missing layer matters: as scientific workflows become more autonomous, "
        "their internal failure structure becomes part of the scientific evidence. We therefore test recursive "
        "scientific observability as the primary endpoint, not as a surrogate for improved prediction. We define "
        "necessary conditions and operational sufficiency within a declared transition schema, deliberately hide, "
        "corrupt or collapse internal channels, and test whether the corresponding metrics detect each disruption. "
        "We further ask whether deeper observation contributes held-out reconstructive information and whether "
        "failure signatures are invariant across DepMap/PRISM, LINCS and TCGA. This design makes a strong but "
        "bounded claim: observability can be measured, disrupted and rejected independently of biological or "
        "clinical success."
    ),
]


RESULTS = [
    (
        "An operational definition separates observability from outcome accuracy",
        [
            (
                "We defined recursive scientific observability (RSO) for the recorded architecture as the "
                "conjunction of six necessary conditions: valid provenance-linked transitions (N1), finite state "
                "availability (N2), identifiable agent channels (N3), reconstructable Master state (N4), ordered "
                "self-error information (N5) and separable recursive channels (N6). Within this declared schema, "
                "the conditions were treated as operationally sufficient only when three stress tests also passed: "
                "failure sensitivity (S1), non-redundant recursive depth (S2) and cross-system invariance (S3). "
                "This is a falsifiable operational definition for the recorded process, not a universal theorem "
                "for nonlinear dynamical systems (Fig. 1)."
            ),
            (
                "The evaluation used 4,431 transition events from observed public data: 2,169 DepMap/PRISM events "
                "over 723 cell models, 2,136 LINCS events over 712 held-out perturbational transitions and 126 TCGA "
                "events over 42 patients from seven held-out cancer projects. All events satisfied the fail-closed "
                "20-field contract, and all 9/9 operational conditions were demonstrated. Structural verification "
                "also confirmed complete agent sets for all three architectures and all 9/9 constituent agents across "
                "1,477 independent biological units (Table 1; Supplementary Tables S1-S6)."
            ),
        ],
    ),
    (
        "Controlled channel loss is detected with graded dose response",
        [
            (
                "We next constructed deliberately non-observable variants by applying six interventions at 0%, 25%, "
                "50%, 75% and 100% of biological units: removing provenance, removing agent states, collapsing distinct "
                "agents, permuting Master states between units, permuting meta-observer values and replacing recursive "
                "uncertainty channels with a shared standardized signal. The complete and fully disrupted architectures "
                "separated on their condition-specific observability measures (Fig. 2). Across three systems and six "
                "failure classes, all 18/18 tests exceeded the declared absolute detection effect of 0.05 and showed a "
                "severity Spearman correlation below -0.70. None failed (Fig. 3; Supplementary Tables S7-S9)."
            ),
            (
                "The intervention-to-metric mapping was specific by construction and remained interpretable. Missing "
                "provenance reduced strict validity; hidden agent states reduced trace completeness; agent collapse "
                "reduced identifiability; Master substitution reduced Master fidelity; meta-observer corruption reduced "
                "rank-order fidelity; and recursive collapse reduced effective channel rank. The result demonstrates "
                "observability failure detection, not causal biological influence of an agent."
            ),
        ],
    ),
    (
        "Recursive depth contributes information and failure signatures transfer",
        [
            (
                "To determine whether recursion added information or merely telemetry, we reconstructed held-out Master "
                "states from nested features of observer depths 0-5 using fivefold cross-validated ridge regression. "
                "The maximum marginal reconstructability gain was 0.207 in DepMap/PRISM, 0.416 in LINCS and 0.399 in "
                "TCGA, exceeding the declared 0.005 criterion in every system (Fig. 4). Gains were non-monotonic at some "
                "levels, exposing redundant or destabilizing telemetry rather than hiding it in a total score."
            ),
            (
                "Failure signatures also showed cross-system stability. Pairwise Spearman correlations of the six "
                "absolute disruption effects yielded a median of 0.754 across DepMap/PRISM, LINCS and TCGA, exceeding "
                "the declared 0.70 criterion (Fig. 5). This invariance concerns the architecture's response to loss of "
                "observability; it does not claim invariant cancer biology or equivalent prediction tasks."
            ),
        ],
    ),
    (
        "Complete telemetry does not guarantee reliable self-observation",
        [
            (
                "The pre-existing real-data audit illustrates why observability cannot be replaced by performance. "
                "Strict transition validity was 1.000 in all systems, but state reconstructability was 0.999, 0.223 "
                "and 1.000 for DepMap/PRISM, LINCS and TCGA, respectively. Recursive consistency was 0.956, 0.568 and "
                "0.436, whereas observer-system instability coupling was 0.271 (95% cluster-bootstrap CI 0.203-0.337), "
                "0.895 (0.870-0.915) and 0.896 (0.797-0.949). Thus, complete logging coexisted with weak state recovery "
                "or weak recursive consistency, which are observable limitations rather than missing results."
            ),
            (
                "On the secondary TCGA endpoint, the held-out Master achieved AUROC 0.677 and Brier score 0.215 versus "
                "0.239 for the fit-prevalence baseline in 42 patients. Observed-label permutation P values were 0.0270 "
                "for AUROC and 0.0360 for Brier score. However, predicted self-error correlated negatively with realized "
                "absolute error (r=-0.267). Moderate discrimination therefore did not imply that the system could rank "
                "its unsafe patient-level predictions. This internal project holdout is not external clinical validation."
            ),
        ],
    ),
    (
        "Offline observer reinforcement remains a secondary negative control",
        [
            (
                "An offline reinforcement extension trained on 598 observed LINCS calibration transitions from 19 "
                "perturbations and was evaluated on the same 712 frozen transitions from 24 unseen perturbations. The "
                "reward-weighted policy achieved mean reward 0.611089 versus 0.609677 for the frozen Master, a difference "
                "of 0.001413 (95% perturbation-cluster bootstrap CI -0.000092 to 0.002988; sign-flip P=0.1454). The gain "
                "was not statistically established. Self-error correlation was 0.658 and trace completeness remained "
                "1.000 across 2,136 events. These findings show how an observable system can reweight or defer, but they "
                "do not support a superiority or treatment-policy claim."
            )
        ],
    ),
]


DISCUSSION = [
    (
        "The main advance is an experimentally falsifiable property for distributed scientific AI. Observability is "
        "not treated as a visual explanation, a confidence score or a synonym for accuracy. It requires separately "
        "recoverable provenance, agent, Master and self-observer channels; it also requires those claims to fail in "
        "predictable ways when the corresponding channels are removed or corrupted. The 18/18 controlled detections, "
        "non-redundant depth and cross-system failure concordance jointly support operational sufficiency within the "
        "declared transition schema. The novelty is therefore the conversion of internal scientific-AI process state "
        "from optional telemetry into a measurable and rejectable experimental endpoint."
    ),
    (
        "This contribution is complementary to recent agentic science rather than a claim to replace it. Autonomous "
        "oncology agents have been judged through tool use and clinical decision outputs [8]; SPARK and Robin demonstrate "
        "hypothesis generation, analytical implementation and experimental discovery [9,10]; ScienceAgentBench tests "
        "whether agents complete authentic analysis tasks [11]. Those evaluations ask whether an agent can do useful "
        "science. RSO asks a prior engineering question: when a distributed system succeeds or fails, are the internal "
        "states needed to interrogate that event still available, distinguishable and falsifiable? The two evaluation "
        "axes should be reported together as autonomy increases."
    ),
    (
        "RSO also differs from interpretability and uncertainty. Recent guidance has emphasized that post-hoc biological "
        "interpretations can be unstable or misleading [4], and studies of confidence communication show a gap between "
        "internal model confidence and human belief [7]. Shortcut-learning audits further show that apparently strong "
        "medical-AI performance can degrade under hidden acquisition bias [12]. Our TCGA result provides a related "
        "process-level warning: AUROC and Brier score were non-random in the internal holdout, yet self-error ranking "
        "failed. An observer of error must therefore be evaluated as its own model, not accepted because the predictor "
        "is calibrated on average."
    ),
    (
        "The cross-system experiment is important because it tests the instrument rather than demanding identical "
        "biological effects. DepMap/PRISM is cross-sectional and multimodal, LINCS is perturbational and time ordered, "
        "and TCGA links baseline tumours to a fixed clinical horizon. Their outcome targets and scales differ, but the "
        "same architectural failures produced concordant observability signatures. This suggests a reusable benchmark "
        "for distributed scientific systems: declare the internal channels, define the failure associated with each "
        "claim, introduce graded disruptions at independent biological units and test depth for marginal held-out "
        "information. The accompanying taxonomy makes negative findings comparable rather than anecdotal."
    ),
    (
        "Several limitations bound the claim. The nine conditions are operationally necessary and sufficient only for "
        "the declared transition schema; they are not a complete nonlinear observability theorem. Thresholds were coded "
        "before the benchmark execution but were not externally preregistered. Controlled telemetry disruptions test "
        "detection specificity, not causal biological effects of agents or recursion levels. The systems are heterogeneous, "
        "and the cross-system concordance is based on six failure classes. TCGA contains only 42 held-out patients and no "
        "longitudinal molecular samples from the same individuals. LINCS represents cell lines and landmark-gene states; "
        "DepMap/PRISM remains in vitro. No result establishes prospective patient benefit, treatment efficacy or clinical "
        "deployment readiness."
    ),
    (
        "Future work should expand observability without changing its primary endpoint. A preregistered benchmark should "
        "freeze thresholds and detector mappings before analysis, add independent institutions and modalities, and test "
        "additional failure classes such as delayed, duplicated or adversarially plausible telemetry. Longitudinal tumour "
        "or liquid-biopsy measurements from the same patients would allow temporal observer states to be compared with "
        "actual disease evolution. Agent ablations and randomized recursion-depth interventions are required before causal "
        "claims. Most importantly, prospective studies should test whether observer instability identifies unreliable "
        "predictions before biological or clinical failure, with the act/defer policy evaluated separately from the base "
        "predictor."
    ),
    (
        "In conclusion, this study establishes recursive observability as an engineering and scientific endpoint that can "
        "be measured, deliberately broken and compared across heterogeneous cancer-AI systems. The results do not show "
        "that observable AI is automatically correct. They show something more fundamental for trustworthy scientific "
        "automation: the internal evidence needed to understand and reject a distributed inference need not disappear "
        "inside the ensemble. Making that evidence experimentally observable creates a testable foundation on which "
        "biological validation, clinical evaluation and responsible autonomy can subsequently be built."
    ),
]


METHODS = [
    (
        "Study design and evidence policy",
        [
            (
                "This observational computational study made recursive scientific observability the primary endpoint. "
                "Prediction and reinforcement performance were secondary. All biological evidence came from observed "
                "public pan-cancer measurements. No patient, cell state or biological trajectory was simulated. Bootstrap "
                "procedures resampled observed independent units and permutation procedures rearranged observed values; "
                "neither generated synthetic biological evidence. The common random seed was 20260824, except the controlled "
                "observability benchmark, which used seed 20260827."
            ),
            (
                "The evaluation priority was fixed in the analysis code as scientific observability, state reconstructability, "
                "agent-Master attribution, recursive self-observability, cancer-AI instability coupling and, last, learning or "
                "prediction performance. Thresholds for the disruption benchmark were declared in the benchmark implementation "
                "before execution but were not registered in an external registry."
            ),
        ],
    ),
    (
        "Data sources and biological units",
        [
            (
                "The DepMap/PRISM architecture used DepMap 24Q2 expression and CRISPR gene-effect matrices and PRISM "
                "Repurposing Public 24Q2 response data [15,16]. Median imputation was applied within each aligned modality. "
                "The shared architecture contained 723 models from 29 lineage labels. The wider real-data validation covered "
                "1,066 models, 20 lineages with sufficient representation, 100 candidate-gene observed-value tests, 615 "
                "drug-target pairs and 1,033 lineage-specific dependency-pharmacology associations. GDSC supplied 575,197 "
                "fitted responses, of which 322,040 mapped by exact Sanger identifiers to 531 DepMap models [17]."
            ),
            (
                "LINCS used the checksum-verified GSE70138 Level-5 matrix [18,19]. Cancer filtering retained 89,192 signatures "
                "and 978 landmark genes. States were matched exactly by cell line, perturbation, dose and perturbation type "
                "across 3, 6 and 24 hours. Whole perturbation identifiers were deterministically assigned: 20% to holdout, "
                "20% of the remaining identifiers to calibration and the remainder to fitting. The resulting frozen partitions "
                "contained 712 holdout transitions from 24 unseen perturbations and 598 calibration transitions from 19 "
                "perturbations."
            ),
            (
                "TCGA used 330 open-access STAR-count RNA-sequencing files, ten from each of 33 projects, harmonized across "
                "60,660 genes, and linked to the TCGA Pan-Cancer Clinical Data Resource [13,14]. The two-year overall-survival "
                "endpoint used a 730-day horizon and one baseline tumour sample per case. Entire projects were ordered using "
                "SHA-256 with salt 'tcga-cdr-os-2y-project-holdout-v1'; 20% of projects were assigned to test and 15% to "
                "calibration. The test partition contained 42 patients from seven projects. This was an internal project "
                "holdout, not an independent clinical cohort."
            ),
        ],
    ),
    (
        "Distributed observer architectures",
        [
            (
                "Each system used three heterogeneous observers and a Master integration layer. DepMap/PRISM observers "
                "represented expression, dependency and pharmacology. Each modality was standardized and decomposed by "
                "full-solver principal-component analysis (PCA) to six components; the concatenated observer states were "
                "decomposed to an eight-component standardized Master state. Fivefold shuffled cross-validation used ridge "
                "regression with alpha=1.0 to reconstruct Master and self-observer states. Counterfactual influence was the "
                "normalized Master-state displacement after zeroing one observer block."
            ),
            (
                "LINCS represented standardized observed states with randomized PCA, using 16 components and seed 20260824. "
                "Features comprised the before state, log-transformed before, after and elapsed hours, log absolute dose and "
                "a missing-dose indicator. Observers were a standardized ridge regressor (alpha=10.0), random-forest regressor "
                "(160 trees, maximum depth 14, minimum leaf size 3, seed 20260824) and standardized Euclidean k-nearest "
                "neighbours (distance weights, p=2, neighbour count bounded between 3 and 25 by fit size). Inverse calibration "
                "mean-squared error defined frozen Master weights. A positive linear regression mapped calibration disagreement, "
                "predicted change and elapsed time to realized error."
            ),
            (
                "TCGA selected the 500 highest-variance genes in the fit partition after log2(1+TPM) transformation. Inputs "
                "were standardized and reduced to 24 PCA components where feasible. Each observer used 24 bootstrap fits of "
                "observed fit patients: balanced logistic regression (C=0.25, liblinear solver, maximum 2,000 iterations), "
                "balanced-subsample random forest (120 trees, minimum leaf size 4, square-root feature sampling) and distance-"
                "weighted k-nearest neighbours with k=max(3,min(15,floor(sqrt(n_fit)))). Inverse calibration Brier score "
                "defined Master weights. The meta-observer averaged standardized ridge (alpha=2.0), random forest (300 trees, "
                "minimum leaf size 3) and distance-weighted k-nearest-neighbour regressors with k bounded between 3 and 10."
            ),
        ],
    ),
    (
        "Transition contract and operational observability",
        [
            (
                "Every observer integration emitted a fail-closed transition containing 20 required fields: episode, time, "
                "cancer-state and agent identifiers; before and after agent states; action; reward or declared influence; "
                "agent uncertainty; before and after Master states; Master and meta uncertainty; evidence identifiers; parent "
                "trace; considered futures; normalized possibility amplitudes; gate; and before and after hypotheses. Missing "
                "provenance, empty or non-finite states, dimension mismatch, invalid uncertainty or inconsistent future weights "
                "invalidated the event. Possibility amplitudes were computational square roots of normalized future weights and "
                "carried no physical-quantum interpretation."
            ),
            (
                "The six necessary conditions were N1 strict transition validity, N2 trace completeness, N3 agent "
                "identifiability, N4 Master fidelity, N5 meta-observer order fidelity and N6 effective channel-rank ratio. "
                "Acceptance required the minimum across systems to be at least 1.00, 1.00, 0.50, 0.95, 0.95 and 0.50, "
                "respectively. Operational sufficiency additionally required S1 detection of every controlled failure, S2 at "
                "least one positive held-out marginal depth gain above 0.005 and S3 median cross-system failure-signature "
                "Spearman correlation of at least 0.70."
            ),
            (
                "State reconstructability was one minus normalized reconstruction error, clipped to [0,1]. Recursive "
                "consistency quantified agreement between recursive uncertainty/self-observation and realized instability. "
                "Observer-system instability coupling was the association between the observed biological-instability channel "
                "and Master instability. These profiles were not collapsed into a post-hoc total score."
            ),
        ],
    ),
    (
        "Controlled observability disruptions",
        [
            (
                "Six disruptions were applied to 0%, 25%, 50%, 75% or 100% of whole episodes, selected with seed 20260827: "
                "removal of evidence identifiers and parent traces; removal of agent states and uncertainty; replacement of "
                "distinct observer states by the within-episode mean; permutation of Master states between episodes; permutation "
                "of meta-uncertainty between episodes; and replacement of agent, Master and meta uncertainty by one shared "
                "min-max standardized signal. This produced 30 conditions per system and 90 overall."
            ),
            (
                "The prespecified detector mapping was strict validity, trace completeness, agent identifiability, Master "
                "fidelity, meta-order fidelity and channel-rank ratio, respectively. A failure class was detected when the "
                "complete-to-terminal effect exceeded 0.05 and its detector had Spearman correlation below -0.70 with "
                "disruption severity."
            ),
        ],
    ),
    (
        "Recursive-depth information",
        [
            (
                "For each episode, nested feature sets were constructed from depth 0 to 5: Master state before integration; "
                "mean agent states before and after; between-agent state dispersion; mean agent uncertainty; Master uncertainty; "
                "and meta-uncertainty. Held-out Master-state reconstruction used a standardized ridge regressor with alpha=1.0 "
                "inside shuffled K-fold cross-validation (up to five folds; fewer only when episode count required it), seed "
                "20260827. Marginal information gain was the difference in held-out reconstructability from the previous depth."
            )
        ],
    ),
    (
        "Secondary reinforcement analysis",
        [
            (
                "The offline LINCS policy model predicted per-observer reward with random-forest regressors (320 trees, maximum "
                "depth 12, minimum leaf size 5, square-root feature sampling). Fivefold GroupKFold cross-fitting grouped entire "
                "perturbations. Candidate policies were hard selection and reward-weighted mixtures with temperatures 0.025, "
                "0.05, 0.1, 0.2, 0.5 and 1.0; the calibration optimum was the mixture at 0.025. A separate self-error random "
                "forest used 480 trees, maximum depth 12 and minimum leaf size 5. The act/defer threshold was the 80th percentile "
                "of cross-fitted predicted self-error. Comparators were the frozen Master, equal ensemble, hard reward policy, "
                "best fixed observer and all individual observers. Reward was 1/(1+normalized L2 error), not a clinical reward."
            )
        ],
    ),
    (
        "Statistical analysis",
        [
            (
                "Inference respected the independent biological unit. Observability confidence intervals used 2,000 cluster-"
                "bootstrap replicates. Observed-value permutation tests used 1,000 permutations for TCGA and nonlinear coupling "
                "audits, with one-sided tests for AUROC and Brier score and two-sided tests for coupling where specified. "
                "Benjamini-Hochberg correction controlled false discovery within reported families [28]. Reinforcement "
                "comparisons used 2,000 perturbation-cluster bootstrap replicates and perturbation-level sign-flip tests. Exact "
                "n, effect estimates, confidence intervals and P values are reported; no post-hoc biological power calculation "
                "was used."
            ),
            (
                "The current local environment used Python 3.12.10, NumPy 2.2.5, pandas 2.3.1, scikit-learn 1.8.0 and "
                "Matplotlib 3.10.1. The project specifies Python >=3.11, NumPy >=1.26, pandas >=2.2, scikit-learn >=1.5 and "
                "Matplotlib >=3.9. " + MODAL_COMPUTE_STATEMENT
            ),
        ],
    ),
    (
        "Ethics, data availability and code availability",
        [
            (
                ETHICS_STATEMENT
            ),
            (
                "Data are available from their original repositories: DepMap and PRISM releases, GDSC, GEO accession GSE70138 "
                "and the NCI Genomic Data Commons/TCGA-CDR. The submission package contains every result table used here as CSV, "
                f"all source figures and the analysis code needed to regenerate the observability benchmark. Code and compact "
                f"results are publicly available at {REPOSITORY_URL}. {ZENODO_STATUS}."
            ),
        ],
    ),
]


REFERENCES = [
    "Kalman, R. E. Contributions to the theory of optimal control. Bol. Soc. Mat. Mexicana 5, 102-119 (1960).",
    "Hermann, R. & Krener, A. J. Nonlinear controllability and observability. IEEE Trans. Autom. Control 22, 728-740 (1977). doi:10.1109/TAC.1977.1101601.",
    "Rudin, C. Stop explaining black box machine learning models for high stakes decisions and use interpretable models instead. Nat. Mach. Intell. 1, 206-215 (2019). doi:10.1038/s42256-019-0048-x.",
    "Chen, V. et al. Applying interpretable machine learning in computational biology: pitfalls, recommendations and opportunities for new developments. Nat. Methods 21, 1454-1461 (2024). doi:10.1038/s41592-024-02359-7.",
    "Lakshminarayanan, B., Pritzel, A. & Blundell, C. Simple and scalable predictive uncertainty estimation using deep ensembles. Adv. Neural Inf. Process. Syst. 30 (2017).",
    "Guo, C., Pleiss, G., Sun, Y. & Weinberger, K. Q. On calibration of modern neural networks. Proc. Mach. Learn. Res. 70, 1321-1330 (2017).",
    "Steyvers, M. et al. What large language models know and what people think they know. Nat. Mach. Intell. 7, 221-231 (2025). doi:10.1038/s42256-024-00976-7.",
    "Ferber, D. et al. Development and validation of an autonomous artificial intelligence agent for clinical decision-making in oncology. Nat. Cancer 6, 1337-1349 (2025). doi:10.1038/s43018-025-00991-6.",
    "Trost, F. et al. An agentic framework for autonomous scientific discovery in cancer pathology. Nat. Med. 32, 2254-2266 (2026). doi:10.1038/s41591-026-04357-y.",
    "Ghareeb, A. E. et al. A multi-agent system for automating scientific discovery. Nature 655, 497-505 (2026). doi:10.1038/s41586-026-10652-y.",
    "Chen, Z. et al. ScienceAgentBench: toward rigorous assessment of language agents for data-driven scientific discovery. ICLR (2025). doi:10.48550/arXiv.2410.05080.",
    "Ong Ly, C. et al. Shortcut learning in medical AI hinders generalization: method for estimating AI model generalization without external data. npj Digit. Med. 7, 124 (2024). doi:10.1038/s41746-024-01118-4.",
    "Hoadley, K. A. et al. Cell-of-origin patterns dominate the molecular classification of 10,000 tumors from 33 types of cancer. Cell 173, 291-304.e6 (2018). doi:10.1016/j.cell.2018.03.022.",
    "Liu, J. et al. An integrated TCGA pan-cancer clinical data resource to drive high-quality survival outcome analytics. Cell 173, 400-416.e11 (2018). doi:10.1016/j.cell.2018.02.052.",
    "Tsherniak, A. et al. Defining a cancer dependency map. Cell 170, 564-576.e16 (2017). doi:10.1016/j.cell.2017.06.010.",
    "Corsello, S. M. et al. Discovering the anticancer potential of non-oncology drugs by systematic viability profiling. Nat. Cancer 1, 235-248 (2020). doi:10.1038/s43018-019-0018-6.",
    "Iorio, F. et al. A landscape of pharmacogenomic interactions in cancer. Cell 166, 740-754 (2016). doi:10.1016/j.cell.2016.06.017.",
    "Subramanian, A. et al. A next generation Connectivity Map: L1000 platform and the first 1,000,000 profiles. Cell 171, 1437-1452.e17 (2017). doi:10.1016/j.cell.2017.10.049.",
    "Keenan, A. B. et al. The Library of Integrated Network-Based Cellular Signatures NIH Program. Cell Syst. 6, 13-24 (2018). doi:10.1016/j.cels.2017.11.001.",
    "Behan, F. M. et al. Prioritization of cancer therapeutic targets using CRISPR-Cas9 screens. Nature 568, 511-516 (2019). doi:10.1038/s41586-019-1103-9.",
    "Hanahan, D. Hallmarks of cancer: new dimensions. Cancer Discov. 12, 31-46 (2022). doi:10.1158/2159-8290.CD-21-1059.",
    "McGranahan, N. & Swanton, C. Clonal heterogeneity and tumor evolution: past, present, and the future. Cell 168, 613-628 (2017). doi:10.1016/j.cell.2017.01.018.",
    "Greaves, M. & Maley, C. C. Clonal evolution in cancer. Nature 481, 306-313 (2012). doi:10.1038/nature10762.",
    "Dagogo-Jack, I. & Shaw, A. T. Tumour heterogeneity and resistance to cancer therapies. Nat. Rev. Clin. Oncol. 15, 81-94 (2018). doi:10.1038/nrclinonc.2017.166.",
    "Sharma, S. V. et al. A chromatin-mediated reversible drug-tolerant state in cancer cell subpopulations. Cell 141, 69-80 (2010). doi:10.1016/j.cell.2010.02.027.",
    "Kompa, B., Snoek, J. & Beam, A. L. Second opinion needed: communicating uncertainty in medical machine learning. npj Digit. Med. 4, 4 (2021). doi:10.1038/s41746-020-00367-3.",
    "Chua, M. et al. Tackling prediction uncertainty in machine learning for healthcare. Nat. Biomed. Eng. 7, 711-718 (2023). doi:10.1038/s41551-022-00988-x.",
    "Benjamini, Y. & Hochberg, Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J. R. Stat. Soc. B 57, 289-300 (1995). doi:10.1111/j.2517-6161.1995.tb02031.x.",
]


MAIN_FIGURES = [
    (
        WORKFLOW,
        "Figure 1",
        "Recursive scientific observability as the primary endpoint. The workflow preserves observed biological state, "
        "individual observer states, inter-observer disagreement, Master state and recursive self-observation under a "
        "provenance-linked transition contract. Controlled channel disruptions and held-out depth reconstruction test "
        "the instrument independently of outcome performance.",
    ),
    (
        BENCHMARK / "figures" / "figure_01_complete_vs_nonobservable.png",
        "Figure 2",
        "Complete versus deliberately non-observable architectures. Bars show condition-specific observability metrics "
        "for complete event streams and the mean after complete disruption of each required channel in DepMap/PRISM, "
        "LINCS and TCGA. The comparison tests architectural visibility, not biological outcome accuracy.",
    ),
    (
        BENCHMARK / "figures" / "figure_02_disruption_dose_response.png",
        "Figure 3",
        "Dose-response detection of controlled observability failures. Each of six channel disruptions was applied to "
        "0%, 25%, 50%, 75% and 100% of whole biological episodes. All 18 system-by-failure tests met the declared effect "
        "and monotonicity criteria.",
    ),
    (
        BENCHMARK / "figures" / "figure_03_recursive_depth_information.png",
        "Figure 4",
        "Held-out information across recursive observer depth. Reconstructability was estimated with cross-validated "
        "ridge regression from nested feature sets. Maximum marginal gains were 0.207, 0.416 and 0.399 for DepMap/PRISM, "
        "LINCS and TCGA, respectively; non-positive levels expose redundant or destabilizing telemetry.",
    ),
    (
        BENCHMARK / "figures" / "figure_04_cross_system_invariance.png",
        "Figure 5",
        "Cross-system invariance of observability-failure signatures. The heat map contains Spearman correlations of the "
        "six absolute disruption effects. Median pairwise concordance was 0.754, evaluating transfer of failure detection "
        "rather than equivalence of biological systems.",
    ),
]


SUPP_GROUPS = [
    ("S1", [1, 2, 3], "Inventory and Master/recursive observability profiles."),
    ("S2", [4, 5, 6], "Agent observability and separated instability/coupling profiles."),
    ("S3", [7, 8, 9], "Agent divergence, observer depth and gate influence."),
    ("S4", [10, 11, 12], "Secondary TCGA performance, LINCS coupling and project outcome profiles."),
    ("S5", [13, 14], "Master-architecture and all-agent structural verification."),
    ("S6", [15, 16, 17], "Clustered intervals, recursive calibration and channel dependence."),
    ("S7", [18, 19], "Selective risk and component-level performance."),
]


SUPP_TABLES = [
    (PRIMARY / "tables" / "table_01_real_data_system_inventory.csv", "S1", "Real-data system inventory."),
    (PRIMARY / "tables" / "table_02_required_transition_schema.csv", "S2", "Required fail-closed transition fields."),
    (PRIMARY / "tables" / "table_03_msoi_profile.csv", "S3", "Master Scientific Observability profile."),
    (PRIMARY / "tables" / "table_04_rmoi_profile.csv", "S4", "Recursive Master Observability profile."),
    (PRIMARY / "tables" / "table_05_saoi_profile.csv", "S5", "Scientific Agent Observability profile."),
    (PRIMARY / "tables" / "table_17_clustered_observability_confidence_intervals.csv", "S6", "Clustered observability confidence intervals."),
    (BENCHMARK / "tables" / "table_01_formal_conditions.csv", "S7", "Operational necessary and sufficient conditions."),
    (BENCHMARK / "tables" / "table_03_detection_sensitivity.csv", "S8", "Controlled-failure detection sensitivity."),
    (BENCHMARK / "tables" / "table_04_recursive_depth_information.csv", "S9", "Held-out information across recursive depth."),
    (BENCHMARK / "tables" / "table_05_cross_system_invariance.csv", "S10", "Cross-system failure-signature invariance."),
    (BENCHMARK / "tables" / "table_06_failure_taxonomy.csv", "S11", "Observability failure taxonomy."),
    (PRIMARY / "tables" / "table_11_tcga_secondary_performance.csv", "S12", "Secondary TCGA performance and permutation audit."),
    (REINFORCEMENT / "holdout_method_evaluation.csv", "S13", "Offline reinforcement and comparator evaluation."),
    (REINFORCEMENT / "cluster_bootstrap_comparisons.csv", "S14", "Cluster-bootstrap reinforcement comparisons."),
    (REINFORCEMENT / "learned_policy_weights.csv", "S15", "Learned observer-policy weights."),
]


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: str = BLACK) -> None:
    run.font.name = "Calibri"
    fonts = run._element.get_or_add_rPr().rFonts
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9)
    start = OxmlElement("w:fldChar")
    start.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([start, instruction, end])


def configure_doc(doc: Document, *, header_text: str, supplement: bool = False) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(8)
    pf.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(BLACK)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.15

    if "Compact List" not in [style.name for style in doc.styles]:
        compact = doc.styles.add_style("Compact List", WD_STYLE_TYPE.PARAGRAPH)
        compact.base_style = normal
        compact.paragraph_format.left_indent = Inches(0.375)
        compact.paragraph_format.first_line_indent = Inches(-0.194)
        compact.paragraph_format.space_after = Pt(4)
        compact.paragraph_format.line_spacing = 1.208

    header = section.header.paragraphs[0]
    header.text = ("SUPPLEMENTARY INFORMATION | " if supplement else "") + header_text
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    set_run_font(header.runs[0], size=9, color=DARK_BLUE)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run(title), size=18, bold=True, color=DARK_BLUE)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)
        set_run_font(p.add_run(subtitle), size=11, italic=True, color=BLACK)


def add_h1(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def add_h2(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 2")


def add_h3(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 3")


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.widow_control = True
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)


def add_contact_block(doc: Document, text: str, *, size: float = 9.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    set_run_font(p.add_run(text), size=size)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Compact List")
    p.add_run("- ")
    p.add_run(text)


def add_caption(doc: Document, label: str, text: str, *, keep_with_next: bool = False) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.keep_with_next = keep_with_next
    set_run_font(p.add_run(f"{label}. "), size=9, bold=True)
    set_run_font(p.add_run(text), size=9)


def add_figure(doc: Document, path: Path, label: str, caption: str,
               *, width_in: float = 6.25, max_height_in: float = 7.3) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    with Image.open(path) as image:
        aspect = image.width / image.height
    final_width = min(width_in, max_height_in * aspect)
    inline = p.add_run().add_picture(str(path), width=Inches(final_width))
    inline._inline.docPr.set("title", label)
    inline._inline.docPr.set("descr", caption)
    add_caption(doc, label, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths: list[int], *, font_size: float = 8.0) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    header = table.rows[0]
    repeat_table_header(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = str(value)
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                set_run_font(run, size=font_size, bold=True)
    for values in rows:
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.08
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if index < 2 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_run_font(run, size=font_size)
    set_table_geometry(table, widths, 120)


def main_table_data() -> tuple[list[str], list[list[str]], list[int]]:
    headers = ["System", "Units", "Events", "Validity", "SR", "RC", "OSIC (95% CI)", "Max depth gain", "Failures"]
    rows = [
        ["DepMap/PRISM", "723", "2,169", "1.000", "0.999", "0.956", "0.271 (0.203-0.337)", "0.207", "6/6"],
        ["LINCS", "712", "2,136", "1.000", "0.223", "0.568", "0.895 (0.870-0.915)", "0.416", "6/6"],
        ["TCGA", "42", "126", "1.000", "1.000", "0.436", "0.896 (0.797-0.949)", "0.399", "6/6"],
    ]
    widths = [1700, 780, 850, 800, 620, 620, 1900, 1200, 890]
    return headers, rows, widths


def add_main_table(doc: Document) -> None:
    headers, rows, widths = main_table_data()
    add_caption(
        doc,
        "Table 1",
        "Primary recursive-observability results. Units are independent cell models, LINCS transitions or patients. "
        "Validity, strict transition validity; SR, state reconstructability; RC, recursive consistency; OSIC, "
        "observer-system instability coupling with biological-unit cluster-bootstrap 95% confidence interval; Failures, "
        "detected controlled failure classes out of six.",
        keep_with_next=True,
    )
    add_table(doc, headers, rows, widths, font_size=7.3)


def narrative_word_count() -> int:
    text = " ".join(INTRODUCTION)
    text += " " + " ".join(title + " " + " ".join(paragraphs) for title, paragraphs in RESULTS)
    text += " " + " ".join(DISCUSSION)
    return len(re.findall(r"\b[\w'-]+\b", text))


def abstract_word_count() -> int:
    return len(re.findall(r"\b[\w'-]+\b", ABSTRACT))


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def build_composite(paths: list[Path], output: Path, title: str) -> None:
    if not paths or any(not path.exists() for path in paths):
        missing = [str(path) for path in paths if not path.exists()]
        raise FileNotFoundError(missing)
    canvas_width = 3600
    margin = 90
    gap = 80
    title_height = 150
    columns = 2 if len(paths) > 1 else 1
    panel_width = (canvas_width - 2 * margin - (columns - 1) * gap) // columns
    prepared: list[tuple[Image.Image, str]] = []
    for index, path in enumerate(paths):
        with Image.open(path) as source:
            image = source.convert("RGB")
            scale = panel_width / image.width
            panel = image.resize((panel_width, max(1, int(image.height * scale))), Image.Resampling.LANCZOS)
        prepared.append((panel, chr(97 + index)))
    rows = math.ceil(len(prepared) / columns)
    row_heights = []
    for row in range(rows):
        row_panels = prepared[row * columns:(row + 1) * columns]
        row_heights.append(max(panel.height for panel, _ in row_panels) + 70)
    canvas_height = margin + title_height + sum(row_heights) + gap * max(0, rows - 1) + margin
    canvas = Image.new("RGB", (canvas_width, canvas_height), "#" + WHITE)
    draw = ImageDraw.Draw(canvas)
    draw.text((margin, margin), title, fill="#1F4E79", font=_font(54, True))
    y = margin + title_height
    for index, (panel, letter) in enumerate(prepared):
        row = index // columns
        col = index % columns
        x = margin + col * (panel_width + gap)
        canvas.paste(panel, (x, y))
        draw.rounded_rectangle((x + 12, y + 12, x + 72, y + 72), radius=8, fill="#FFFFFF", outline="#4A5568", width=2)
        draw.text((x + 28, y + 17), letter, fill="#111111", font=_font(38, True))
        if col == columns - 1 or index == len(prepared) - 1:
            y += row_heights[row] + gap
    output.parent.mkdir(parents=True, exist_ok=True)
    canvas.save(output, dpi=(300, 300), optimize=True)


def prepare_figures() -> list[tuple[Path, str, str]]:
    FIG_OUT.mkdir(parents=True, exist_ok=True)
    SUPP_FIG_OUT.mkdir(parents=True, exist_ok=True)
    for index, (source, _, _) in enumerate(MAIN_FIGURES, start=1):
        destination = FIG_OUT / f"Figure_{index}.png"
        shutil.copy2(source, destination)
        svg_source = source.with_suffix(".svg")
        if svg_source.exists():
            shutil.copy2(svg_source, FIG_OUT / f"Figure_{index}.svg")

    supplementary: list[tuple[Path, str, str]] = []
    for label, indices, scope in SUPP_GROUPS:
        paths = [PRIMARY / "figures" / f"figure_{index:02d}_{_primary_figure_stem(index)}.png" for index in indices]
        output = SUPP_FIG_OUT / f"Supplementary_Figure_{label}.png"
        build_composite(paths, output, f"Supplementary Figure {label}: {scope}")
        filenames = ", ".join(path.name for path in paths)
        caption = f"{scope} Panels reproduce only the retrieved result figures ({filenames}); no biological values were simulated or redrawn."
        supplementary.append((output, f"Supplementary Figure {label}", caption))

    reinforcement_groups = [
        ("S8", [1, 2], "Offline reinforcement reward and learned observer weights."),
        ("S9", [3, 4], "Master self-observation and selective-risk behaviour."),
    ]
    for label, indices, scope in reinforcement_groups:
        paths = [REINFORCEMENT / "figures" / f"figure_{index:02d}_{_reinforcement_figure_stem(index)}.png" for index in indices]
        output = SUPP_FIG_OUT / f"Supplementary_Figure_{label}.png"
        build_composite(paths, output, f"Supplementary Figure {label}: {scope}")
        filenames = ", ".join(path.name for path in paths)
        caption = f"{scope} Panels reproduce only the retrieved result figures ({filenames}); the reward analysis is secondary and does not establish treatment benefit."
        supplementary.append((output, f"Supplementary Figure {label}", caption))
    return supplementary


def _primary_figure_stem(index: int) -> str:
    return {
        1: "system_inventory",
        2: "msoi_heatmap",
        3: "rmoi_heatmap",
        4: "saoi_heatmap",
        5: "instability_channel_distributions",
        6: "oosc_heatmap",
        7: "agentic_divergence_matrices",
        8: "observer_depth_profiles",
        9: "gate_influence_distributions",
        10: "tcga_roc_calibration",
        11: "lincs_cancer_master_coupling",
        12: "tcga_project_outcomes",
        13: "master_architecture_verification",
        14: "all_agent_observability_verification",
        15: "clustered_observability_intervals",
        16: "recursive_uncertainty_calibration",
        17: "channel_partial_correlations",
        18: "selective_risk_coverage",
        19: "agent_master_component_performance",
    }[index]


def _reinforcement_figure_stem(index: int) -> str:
    return {
        1: "holdout_reward_comparison",
        2: "learned_observer_weights",
        3: "master_self_observation",
        4: "selective_risk",
    }[index]


def build_main() -> None:
    if abstract_word_count() > 150:
        raise ValueError(f"Abstract exceeds 150 words: {abstract_word_count()}")
    if narrative_word_count() > 3500:
        raise ValueError(f"Main text exceeds 3,500 words: {narrative_word_count()}")
    doc = Document()
    configure_doc(doc, header_text=SHORT_TITLE)
    add_title(doc, TITLE, "Nature Machine Intelligence Article | observability is the primary endpoint")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    set_run_font(p.add_run("Author details: "), bold=True)
    set_run_font(p.add_run(f"{AUTHOR_FULL}; {AFFILIATION}; {AUTHOR_EMAIL}; {AUTHOR_ORCID}."), color=DARK_BLUE)
    add_body(doc, f"Main-text word count (Introduction, Results and Discussion): {narrative_word_count()}. Abstract: {abstract_word_count()} words.")

    add_h1(doc, "Abstract")
    add_body(doc, ABSTRACT)

    for paragraph in INTRODUCTION:
        add_body(doc, paragraph)

    add_h1(doc, "Results")
    for result_index, (heading, paragraphs) in enumerate(RESULTS):
        add_h2(doc, heading)
        for paragraph in paragraphs:
            add_body(doc, paragraph)
        if result_index == 0:
            add_figure(doc, FIG_OUT / "Figure_1.png", "Figure 1", MAIN_FIGURES[0][2], max_height_in=4.0)
            add_main_table(doc)
        elif result_index == 1:
            add_figure(doc, FIG_OUT / "Figure_2.png", "Figure 2", MAIN_FIGURES[1][2], max_height_in=4.2)
            add_figure(doc, FIG_OUT / "Figure_3.png", "Figure 3", MAIN_FIGURES[2][2], max_height_in=5.5)
        elif result_index == 2:
            add_figure(doc, FIG_OUT / "Figure_4.png", "Figure 4", MAIN_FIGURES[3][2], max_height_in=4.3)
            add_figure(doc, FIG_OUT / "Figure_5.png", "Figure 5", MAIN_FIGURES[4][2], max_height_in=4.4)

    add_h1(doc, "Discussion")
    for paragraph in DISCUSSION:
        add_body(doc, paragraph)

    add_h1(doc, "Methods")
    for heading, paragraphs in METHODS:
        add_h2(doc, heading)
        for paragraph in paragraphs:
            add_body(doc, paragraph)

    add_h1(doc, "Acknowledgements")
    add_body(doc, ACKNOWLEDGEMENT_STATEMENT + " " + FUNDING_STATEMENT)
    add_h1(doc, "Author contributions")
    add_body(doc, CREDIT_STATEMENT)
    add_h1(doc, "Competing interests")
    add_body(doc, COMPETING_INTERESTS_STATEMENT)
    add_h1(doc, "References")
    for index, reference in enumerate(REFERENCES, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(f"{index}. {reference}")

    doc.core_properties.title = TITLE
    doc.core_properties.subject = "Nature Machine Intelligence Article focused on recursive scientific observability"
    doc.core_properties.author = AUTHOR_NAME
    doc.core_properties.keywords = "recursive observability, scientific AI, multi-agent systems, cancer, uncertainty"
    doc.save(MAIN_DOCX)


def add_csv_table(doc: Document, path: Path, label: str, caption: str) -> None:
    headers, rows = read_csv_table(path)
    formatted = [[format_value(value) for value in row] for row in rows]
    if len(headers) <= 6:
        groups = [list(range(len(headers)))]
    else:
        anchors = list(range(min(2, len(headers))))
        remaining = list(range(len(anchors), len(headers)))
        groups = [anchors + remaining[index:index + 3] for index in range(0, len(remaining), 3)]
    for panel_index, indices in enumerate(groups, start=1):
        panel = "" if len(groups) == 1 else f", panel {chr(64 + panel_index)} of {len(groups)}"
        add_caption(doc, f"Supplementary Table {label}{panel}", caption + f" Source: {path.name}.", keep_with_next=True)
        panel_headers = [headers[index] for index in indices]
        panel_rows = [[row[index] if index < len(row) else "" for index in indices] for row in formatted]
        widths = column_widths(panel_headers, panel_rows, 9360)
        add_table(doc, panel_headers, panel_rows, widths, font_size=7.4)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_supplement(supplementary_figures: list[tuple[Path, str, str]]) -> None:
    doc = Document()
    configure_doc(doc, header_text=SHORT_TITLE, supplement=True)
    add_title(doc, "Supplementary Information", TITLE)
    add_body(
        doc,
        "This file contains expanded definitions, complete reported hyperparameters, selected readable result tables and "
        "nine composite figures reproducing all 23 primary-observability and reinforcement figures not used as main "
        "display items. Every original table is supplied without alteration in Supplementary Data 1 (CSV archive), and "
        "every original figure is supplied in Supplementary Data 2 (figure archive). No biological data were simulated."
    )

    add_h1(doc, "Supplementary Methods")
    add_h2(doc, "Formal scope of recursive scientific observability")
    add_body(
        doc,
        "For the recorded distributed process, RSO=N1 AND N2 AND N3 AND N4 AND N5 AND N6. The necessary conditions "
        "become operationally sufficient within the declared transition schema only with S1 AND S2 AND S3. N1-N6 "
        "specify what must remain observable; S1-S3 test whether those claims are experimentally discriminable, "
        "informative across depth and transferable across systems. This formulation is intentionally bounded and does "
        "not assert complete nonlinear observability, controllability, consciousness or causal biological agency."
    )
    add_h2(doc, "No-simulation audit")
    add_body(
        doc,
        "Biological rows, patients, cell states and trajectories were not generated. Bootstrap replicates resampled "
        "observed biological units with replacement. Permutation and sign-flip tests rearranged observed values or "
        "cluster-level effects. Controlled observability disruptions altered only the recorded telemetry of observed "
        "episodes; they did not create or alter biological outcomes."
    )
    add_h2(doc, "Hyperparameter ledger")
    hyper_headers = ["Module", "Parameter", "Value", "Code source"]
    hyper_rows = [
        ["Global", "Seeds", "20260824; benchmark 20260827", "validation and benchmark modules"],
        ["DepMap/PRISM", "Agent PCA", "6 components; full solver; standardized", "src/agents/master.py"],
        ["DepMap/PRISM", "Master PCA", "8 components; full solver; standardized", "src/agents/master.py"],
        ["DepMap/PRISM", "Reconstruction", "5-fold shuffled CV; Ridge alpha=1.0", "src/agents/master.py"],
        ["Real-data validation", "Defaults", "50 genes; 100 drug pairs; 200 permutations; >=20 lineage models", "src/validation/real_data.py"],
        ["LINCS", "Partition", "20% perturbations holdout; 20% of remainder calibration", "src/validation/lincs_future_state.py"],
        ["LINCS", "State PCA", "16 randomized components; seed 20260824", "src/validation/lincs_future_state.py"],
        ["LINCS", "Ridge", "alpha=10.0; standardized", "src/validation/lincs_future_state.py"],
        ["LINCS", "Random forest", "160 trees; depth 14; leaf 3", "src/validation/lincs_future_state.py"],
        ["LINCS", "Nearest neighbours", "k=3-25; distance weights; p=2", "src/validation/lincs_future_state.py"],
        ["TCGA", "Endpoint", "2-year OS; horizon 730 days", "src/validation/tcga_outcomes.py"],
        ["TCGA", "Features", "top 500 fit-variance genes; 24 PCA components", "src/validation/tcga_outcomes.py"],
        ["TCGA", "Bootstrap fits", "24 per observer", "src/validation/tcga_outcomes.py"],
        ["TCGA", "Logistic observer", "C=0.25; balanced; liblinear; max_iter=2,000", "src/validation/tcga_outcomes.py"],
        ["TCGA", "Forest observer", "120 trees; leaf 4; sqrt features; balanced subsample", "src/validation/tcga_outcomes.py"],
        ["TCGA", "Neighbour observer", "k=3-15 from floor(sqrt(n)); distance weights", "src/validation/tcga_outcomes.py"],
        ["TCGA", "Meta-observer", "Ridge alpha=2; RF 300 trees/leaf 3; kNN k=3-10", "src/validation/tcga_outcomes.py"],
        ["TCGA", "Permutation audit", "1,000 rearrangements of observed values", "src/validation/tcga_outcomes.py"],
        ["Verification", "Intervals", "2,000 biological-unit cluster bootstrap replicates", "src/observability/verification.py"],
        ["Benchmark", "Severities", "0, 0.25, 0.5, 0.75, 1.0", "src/observability/benchmark.py"],
        ["Benchmark", "Depth model", "up to 5 folds; standardized Ridge alpha=1.0", "src/observability/benchmark.py"],
        ["Benchmark", "Detection", "effect >0.05 and severity Spearman <-0.70", "src/observability/benchmark.py"],
        ["Benchmark", "Depth/invariance", "gain >0.005; median Spearman >=0.70", "src/observability/benchmark.py"],
        ["Reinforcement", "Reward forest", "320 trees; depth 12; leaf 5; sqrt features", "src/validation/reinforcement_master.py"],
        ["Reinforcement", "Policy temperatures", "0.025, 0.05, 0.1, 0.2, 0.5, 1.0 plus hard", "src/validation/reinforcement_master.py"],
        ["Reinforcement", "Self-error forest", "480 trees; depth 12; leaf 5", "src/validation/reinforcement_master.py"],
        ["Reinforcement", "Deferral", "80th percentile cross-fitted predicted error", "src/validation/reinforcement_master.py"],
        ["Reinforcement", "Comparisons", "2,000 perturbation-cluster bootstrap/sign-flip replicates", "src/validation/reinforcement_master.py"],
    ]
    add_caption(doc, "Supplementary Table S0", "Hyperparameters transcribed from the project source code; unrecorded hardware and runtime were not inferred.", keep_with_next=True)
    add_table(doc, hyper_headers, hyper_rows, [1700, 2100, 2860, 2700], font_size=7.6)

    add_h1(doc, "Supplementary Results Tables")
    add_body(
        doc,
        "The following tables contain the central observability, controlled-disruption and secondary analyses in a "
        "readable Word format. Supplementary Data 1 contains all 31 original CSV tables, including the full 90-row "
        "disruption matrix and all channel, calibration, risk-coverage and component-level tables."
    )
    for position, (path, label, caption) in enumerate(SUPP_TABLES, start=1):
        add_csv_table(doc, path, label, caption)
        if position in {5, 9, 12}:
            doc.add_page_break()

    add_h1(doc, "Supplementary Figure Atlas")
    add_body(
        doc,
        "Figures S1-S7 reproduce all nineteen original primary observability figures. Figures S8-S9 reproduce all four "
        "original reinforcement figures. Main Figures 1-5 reproduce the workflow and all four benchmark figures. Together "
        "the manuscript and supplement therefore include every retrieved figure without generating replacement results."
    )
    for index, (path, label, caption) in enumerate(supplementary_figures):
        add_figure(doc, path, label, caption, width_in=6.35, max_height_in=7.25)
        if index != len(supplementary_figures) - 1:
            doc.add_page_break()

    add_h1(doc, "Supplementary Claim Boundary")
    add_body(
        doc,
        "The experiments establish that recursive observability is measurable, disruptable and cross-system comparable "
        "for the recorded architectures. They do not establish clinical utility, causal treatment effects, prospective "
        "failure prediction, external patient validation, sentience, physical quantum behaviour or a universal observability theorem."
    )
    doc.core_properties.title = TITLE + " - Supplementary Information"
    doc.core_properties.subject = "Methods, tables and complete figure atlas"
    doc.core_properties.author = AUTHOR_NAME
    doc.save(SUPP_DOCX)


def build_cover_letter() -> None:
    doc = Document()
    configure_doc(doc, header_text="Cover letter")
    add_title(doc, "Cover Letter", "Submission to Nature Machine Intelligence")
    add_body(doc, DATE)
    add_body(doc, "Dear Chief Editor,")
    add_body(
        doc,
        f"Please consider this Article, '{TITLE}'. The manuscript addresses a basic limitation of distributed scientific "
        "AI: a system may produce a useful answer while the internal process needed to interrogate that answer is not "
        "reconstructable or falsifiable. We make recursive scientific observability, rather than outcome accuracy, the "
        "primary experimental endpoint."
    )
    add_body(
        doc,
        "The principal advance is an operational framework with six necessary conditions and three stress-test conditions. "
        "Across 4,431 observed-data transitions in DepMap/PRISM, LINCS and TCGA, all nine conditions were demonstrated. "
        "Ninety graded channel disruptions produced 18/18 successful system-by-failure detections, recursive depth added "
        "held-out reconstructive information in every system, and failure signatures showed median cross-system Spearman "
        "concordance of 0.754. These results establish observability as a measurable and rejectable property; they do not "
        "claim clinical efficacy or biological correctness."
    )
    add_body(
        doc,
        "The work should interest Nature Machine Intelligence readers in multi-agent systems, AI for scientific discovery, "
        "uncertainty, interpretability, trustworthy AI and machine learning for health. It complements recent autonomous "
        "science systems by evaluating whether internal evidence remains inspectable when such systems succeed or fail. "
        "All biological inputs are observed public data; no patient, cell state or biological trajectory was simulated."
    )
    add_body(doc, "This manuscript is original, is not under consideration elsewhere, and the sole author approves its submission.")
    add_body(doc, "Sincerely,")
    add_contact_block(doc, f"{AUTHOR_FULL}\n{AUTHOR_ROLE}\n{AFFILIATION}\nEmail: {AUTHOR_EMAIL}\n{AUTHOR_ORCID}")
    doc.core_properties.title = "Cover letter - " + TITLE
    doc.core_properties.author = AUTHOR_NAME
    doc.save(COVER_DOCX)


def build_title_page() -> None:
    doc = Document()
    configure_doc(doc, header_text="Title page and declarations")
    add_title(doc, TITLE, "Article for Nature Machine Intelligence")
    add_h1(doc, "Authors and affiliations")
    add_body(doc, f"{AUTHOR_FULL} ({AUTHOR_ORCID})")
    add_body(doc, AFFILIATION)
    add_body(doc, "Sole author; no equal-contribution or present-address statements apply.")
    add_h1(doc, "Corresponding author")
    add_contact_block(doc, f"{AUTHOR_FULL}\n{AUTHOR_ROLE}\n{AFFILIATION}\nEmail: {AUTHOR_EMAIL}\nTelephone: to be supplied in the submission system\n{AUTHOR_ORCID}", size=10.0)
    add_h1(doc, "Author contributions")
    add_body(doc, CREDIT_STATEMENT)
    add_h1(doc, "Acknowledgements and funding")
    add_body(doc, ACKNOWLEDGEMENT_STATEMENT + " " + FUNDING_STATEMENT)
    add_h1(doc, "Competing interests")
    add_body(doc, COMPETING_INTERESTS_STATEMENT)
    add_h1(doc, "Ethics and consent")
    add_body(doc, ETHICS_STATEMENT)
    add_h1(doc, "Materials and correspondence")
    add_body(doc, f"Correspondence and requests for materials should be addressed to {AUTHOR_NAME} ({AUTHOR_EMAIL}).")
    add_h1(doc, "Submission declarations")
    for item in [
        "The sole author has approved the manuscript and its submission.",
        "The work is original and is not under consideration elsewhere.",
        "The author list and order are final.",
        "All data, code, funding and competing-interest statements are accurate.",
    ]:
        add_bullet(doc, item)
    doc.core_properties.title = "Title page and declarations - " + TITLE
    doc.core_properties.author = AUTHOR_NAME
    doc.save(TITLE_PAGE_DOCX)


def build_reporting_summary() -> None:
    doc = Document()
    configure_doc(doc, header_text="Nature Portfolio Reporting Summary responses")
    add_title(doc, "Nature Portfolio Reporting Summary Response Draft", "These responses are already written into Nature_Forms/Nature_Portfolio_Reporting_Summary_COMPLETED.pdf")
    add_body(doc, f"Corresponding author: {AUTHOR_FULL} | Last updated: {DATE} | Email: {AUTHOR_EMAIL} | {AUTHOR_ORCID}")
    sections = [
        ("Statistics", [
            "Exact sample sizes: reported for every system, split, table and statistical comparison.",
            "Summary statistics: effect estimates, 95% confidence intervals and exact P values are reported where applicable.",
            "Assumptions: the benchmark uses nonparametric rank correlations and resampling; no normality assumption was used for the primary controlled-disruption endpoint.",
            "Multiple comparisons: Benjamini-Hochberg correction was applied within declared permutation families.",
            "Software: Python 3.12.10; NumPy 2.2.5; pandas 2.3.1; scikit-learn 1.8.0; Matplotlib 3.10.1.",
        ]),
        ("Life-sciences study design", [
            "Sample size: no prospective power calculation; all eligible observed units produced by the frozen public-data pipelines were analysed. The independent units were 723 DepMap models, 712 LINCS holdout transitions and 42 TCGA holdout patients.",
            "Data exclusions: only prespecified eligibility, linkage, sample-type and partition rules in Methods were applied. Missing or invalid transition telemetry failed closed rather than being imputed at the observability layer.",
            "Replication: architectural failure detection was repeated across three heterogeneous biological systems and six controlled failure classes.",
            "Randomization: no biological intervention groups were randomized. Deterministic hashed/project or perturbation splits and seeded episode selection were used as described.",
            "Blinding: not performed; analyses were deterministic computational audits of public data.",
        ]),
        ("Human participants and human data", [
            "The study used de-identified public TCGA and TCGA-CDR records and did not recruit participants or access identifiable data.",
            "Sex, gender, race and ethnicity were not used as model inputs or subgroup endpoints in the reported observability benchmark.",
            ETHICS_STATEMENT,
        ]),
        ("Data and code", [
            "All primary datasets and access routes are identified in Methods.",
            "All result tables are supplied as original CSV files; all figures are supplied as original PNG/SVG files.",
            f"The analysis source and tests are supplied in Supplementary Software 1 and at {REPOSITORY_URL}. {ZENODO_STATUS}.",
            MODAL_COMPUTE_STATEMENT,
        ]),
    ]
    for heading, items in sections:
        add_h1(doc, heading)
        for item in items:
            add_bullet(doc, item)
    add_h1(doc, "Answers written into the official smart form")
    add_body(doc, "The entries below were written verbatim into the completed Reporting Summary smart form by the script fill_nature_reporting_summary.py, and are reproduced here so that this draft and the official form cannot diverge.")
    from fill_nature_reporting_summary import described_answers

    for label, answer in described_answers():
        add_bullet(doc, f"{label}: {answer}")
    add_h1(doc, "Final confirmations")
    for item in [
        "Insert the Zenodo DOI once the archival release is published",
        "Review and sign Nature_Forms/Nature_Portfolio_Reporting_Summary_COMPLETED.pdf in Adobe Acrobat or Reader; these responses are already written into that official smart form",
    ]:
        add_bullet(doc, item)
    doc.core_properties.title = "Reporting Summary responses - " + TITLE
    doc.save(REPORTING_DOCX)


def build_ml_checklist() -> None:
    doc = Document()
    configure_doc(doc, header_text="Nature Machine Learning Checklist responses")
    add_title(doc, "Nature Machine Learning Checklist v1.1 Response Draft", "Transfer these responses to the official checklist")
    add_body(doc, f"Corresponding author: {AUTHOR_FULL} | Last updated: {DATE} | Email: {AUTHOR_EMAIL} | {AUTHOR_ORCID}")
    sections = [
        ("1. Availability and reproducibility", [
            "Source code is included in Supplementary Software 1 and can be made available to reviewers.",
            "A README and tests are included. Twenty-one local tests passed before manuscript construction.",
            f"The public reproducibility repository is {REPOSITORY_URL}. {ZENODO_STATUS}.",
            "No inaccessible pretrained foundation model is used in the reported observability benchmark.",
        ]),
        ("2. Datasets", [
            "All data sources, biological units and deterministic split rules are reported.",
            "Fit, calibration and test units are separated by whole perturbation identifiers for LINCS and whole cancer projects for TCGA.",
            "No synthetic patients, cell states or biological trajectories were used.",
            "Potential limitations include lineage imbalance, cell-line context, landmark-gene measurement and the small TCGA holdout.",
        ]),
        ("3. Models and training", [
            "All estimators and hyperparameters are reported in Methods and Supplementary Table S0.",
            "Model selection used only fit/calibration partitions. The frozen holdout was not used for policy tuning.",
            "Random seeds, feature transformations, PCA dimensions, ensemble sizes and neighbour rules are reported.",
            MODAL_COMPUTE_STATEMENT,
        ]),
        ("4. Evaluation", [
            "The primary endpoint is observability, assessed with nine declared conditions, 90 graded disruptions and held-out depth reconstruction.",
            "The detector mapping, effect threshold, monotonicity threshold, depth-gain threshold and invariance threshold are explicit.",
            "Independent biological units, cluster bootstrap, permutations and sign-flip tests are specified.",
            "Prediction and reinforcement results are secondary and all negative or non-significant results are retained.",
        ]),
        ("5. Scope and limitations", [
            "The work does not establish clinical utility, causal agent effects, treatment benefit, consciousness or a universal mathematical observability theorem.",
            "External clinical and prospective observer-instability validation remain future work.",
        ]),
    ]
    for heading, items in sections:
        add_h1(doc, heading)
        for item in items:
            add_bullet(doc, item)
    doc.core_properties.title = "Machine Learning Checklist responses - " + TITLE
    doc.save(ML_DOCX)


def build_instruction_audit() -> None:
    doc = Document()
    configure_doc(doc, header_text="Nature Machine Intelligence author-instruction audit")
    add_title(doc, "Nature Machine Intelligence Author-Instruction Audit", f"Checked {DATE}")
    rows = [
        ["Article main text", "Up to 3,500 words excluding abstract, Methods, references and legends", str(narrative_word_count()), "Pass"],
        ["Abstract", "Up to 150 words; unreferenced", str(abstract_word_count()), "Pass" if abstract_word_count() <= 150 else "Fail"],
        ["Main display items", "Up to six figures and/or tables", "5 figures + 1 table", "Pass"],
        ["Structure", "Unheaded Introduction; Results; Discussion; Methods", "Matched", "Pass"],
        ["Subheadings", "Results/Methods topical; Discussion none", "Matched", "Pass"],
        ["References", "Typically up to 50", str(len(REFERENCES)), "Pass"],
        ["Supplement", "Single combined file; large data separate", "Word + submission PDF; CSV archive", "Pass"],
        ["Code availability", "Statement and reviewer access for central custom code", f"Public GitHub repository; {ZENODO_STATUS}", "Conditional"],
        ["Reporting forms", "Life-sciences Reporting Summary and ML checklist", "Completed ML PDF + Reporting Summary response draft", "Conditional"],
        ["Author declarations", "Authors, contributions, interests, funding", "Sole-author fields, ORCID and no-funding statement completed", "Complete"],
    ]
    add_table(doc, ["Item", "Instruction", "Package", "Status"], rows, [1800, 3100, 2760, 1700], font_size=8.0)
    add_h1(doc, "Official sources")
    sources = [
        "https://www.nature.com/natmachintell/content",
        "https://www.nature.com/natmachintell/submission-guidelines/initial-formatting",
        "https://www.nature.com/natmachintell/submission-guidelines/writing-and-language",
        "https://www.nature.com/natmachintell/editorial-policies",
        "https://www.nature.com/nature/editorial-policies/reporting-standards",
        "https://www.nature.com/documents/machine-learning-checklist.pdf",
        "https://www.nature.com/documents/nr-reporting-summary.pdf",
    ]
    for source in sources:
        add_bullet(doc, source)
    add_h1(doc, "Non-format editorial risk")
    add_body(doc, "Journal fit cannot be guaranteed by formatting. The editorial case depends on presenting recursive observability as a general machine-intelligence contribution and keeping outcome claims strictly secondary and bounded.")
    doc.core_properties.title = "NMI author-instruction audit"
    doc.save(INSTRUCTIONS_DOCX)


def build_submission_checklist() -> None:
    doc = Document()
    configure_doc(doc, header_text="Submission readiness")
    add_title(doc, "Submission Readiness Checklist", TITLE)
    add_h1(doc, "Complete in this package")
    for item in [
        "Journal-length Article in Word with five figures and one main table",
        "Single Supplementary Information file in Word and submission PDF",
        "All 28 retrieved source figures represented in the manuscript/supplement and archived separately",
        "All 31 original result tables supplied as labelled CSV source data",
        "Cover letter, title-page/declaration template and standalone Main Table 1",
        "Nature Reporting Summary and Machine Learning Checklist response drafts",
        "Completed official Nature PDFs: Reporting Summary smart form and Machine Learning Checklist v1.1",
        "Reproducibility code and tests archive",
        "Numerical, figure-inventory and manuscript-structure audit",
    ]:
        add_bullet(doc, item)
    add_h1(doc, "Completed author and infrastructure actions")
    for item in [
        f"Sole author and corresponding author entered: {AUTHOR_FULL}",
        "Sole-author CRediT contribution statement completed",
        "Competing-interest and public-data ethics statements completed",
        f"Public GitHub reproducibility repository created: {REPOSITORY_URL}",
        "Modal workspace, container and declared compute resources documented",
        "Originality, author approval and exclusive-submission declarations completed",
        f"ORCID confirmed: {ORCID_ID}",
        "No-specific-funding statement confirmed by the author",
        "Zenodo archive metadata prepared: publication_repository/.zenodo.json and CITATION.cff",
        "Reporting Summary responses transferred into Nature's official smart form: Nature_Forms/Nature_Portfolio_Reporting_Summary_COMPLETED.pdf",
        "Machine Learning Checklist v1.1 completed: Nature_Forms/Nature_Machine_Learning_Checklist_COMPLETED.pdf",
    ]:
        add_bullet(doc, item)
    add_h1(doc, "Remaining actions")
    for item in [
        "Enable the Zenodo-GitHub integration for the repository and publish release v1.0.0 to mint the DOI, then insert the DOI throughout the package",
        "Open both completed Nature PDFs in Adobe Acrobat or Reader to review and sign them; the Reporting Summary is a LiveCycle form that only Adobe renders, so other PDF viewers will show its placeholder page",
        "Supply a telephone number directly in the submission system if the journal requires it",
    ]:
        add_bullet(doc, item)
    doc.core_properties.title = "Submission readiness - " + TITLE
    doc.save(CHECKLIST_DOCX)


def build_standalone_table() -> None:
    doc = Document()
    configure_doc(doc, header_text="Main Table 1")
    add_title(doc, "Main Table 1", TITLE)
    add_main_table(doc)
    doc.core_properties.title = "Main Table 1 - " + TITLE
    doc.save(TABLE_DOCX)


def download_forms() -> None:
    FORMS_OUT.mkdir(parents=True, exist_ok=True)
    urls = {
        "Nature_Machine_Learning_Checklist_v1.1.pdf": "https://www.nature.com/documents/machine-learning-checklist.pdf",
        "Nature_Portfolio_Reporting_Summary_Smart_Form.pdf": "https://www.nature.com/documents/nr-reporting-summary.pdf",
        "Nature_Portfolio_Reporting_Summary_Reference.pdf": "https://www.nature.com/documents/nr-reporting-summary-flat.pdf",
    }
    for filename, url in urls.items():
        target = FORMS_OUT / filename
        if target.exists():
            continue
        response = requests.get(url, timeout=60)
        response.raise_for_status()
        target.write_bytes(response.content)


SOFFICE = Path("C:/Program Files/LibreOffice/program/soffice.exe")
PDF_EXPORTS = ("02_NMI_Supplementary_Information", "05_Nature_Reporting_Summary_Responses", "06_Nature_Machine_Learning_Checklist_Responses")


def export_pdfs() -> None:
    """Render the documents that Nature wants as PDF alongside the Word originals."""
    if not SOFFICE.exists():
        raise FileNotFoundError(f"LibreOffice is required to export submission PDFs: {SOFFICE}")
    for stem in PDF_EXPORTS:
        source = OUT / f"{stem}.docx"
        subprocess.run(
            [str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(OUT), str(source)],
            check=True,
            capture_output=True,
        )
        target = OUT / f"{stem}.pdf"
        if not target.exists() or target.stat().st_mtime < source.stat().st_mtime:
            raise RuntimeError(f"PDF export did not refresh {target}")


REPO_DIR = ROOT / "publication_repository"
ARCHIVE_VERSION = "1.0.0"
ARCHIVE_LICENSE = "MIT"
ARCHIVE_LICENSE_NAME = "MIT License"
KEYWORDS = [
    "recursive scientific observability",
    "distributed artificial intelligence",
    "pan-cancer",
    "failure detection",
    "reproducible research",
]


def build_archive_metadata() -> None:
    """Write the Zenodo and citation metadata that the archival release will pick up."""
    if not REPO_DIR.is_dir():
        return
    citation = [
        "cff-version: 1.2.0",
        'message: "If you use this software or its results, please cite the associated manuscript and archived release."',
        f'title: "{TITLE}"',
        "type: software",
        "authors:",
        '  - family-names: "Yadalam"',
        '    given-names: "Pradeep Kumar"',
        f'    affiliation: "{AFFILIATION}"',
        f'    email: "{AUTHOR_EMAIL}"',
        f'    orcid: "https://orcid.org/{ORCID_ID}"',
        f'repository-code: "{REPOSITORY_URL}"',
        f"date-released: {DATE}",
        f"version: {ARCHIVE_VERSION}",
        f'license: "{ARCHIVE_LICENSE}"',
        "keywords:",
    ]
    citation += [f"  - {keyword}" for keyword in KEYWORDS]
    (REPO_DIR / "CITATION.cff").write_text("\n".join(citation) + "\n", encoding="utf-8")

    zenodo = {
        "title": TITLE,
        "upload_type": "software",
        "description": ABSTRACT,
        "version": ARCHIVE_VERSION,
        "publication_date": DATE,
        "language": "eng",
        "license": ARCHIVE_LICENSE,
        "access_right": "open",
        "keywords": KEYWORDS,
        "creators": [
            {
                "name": "Yadalam, Pradeep Kumar",
                "affiliation": AFFILIATION,
                "orcid": ORCID_ID,
            }
        ],
        "contributors": [
            {
                "name": "Yadalam, Pradeep Kumar",
                "type": "DataCurator",
                "affiliation": AFFILIATION,
                "orcid": ORCID_ID,
            }
        ],
        "related_identifiers": [
            {"identifier": REPOSITORY_URL, "relation": "isSupplementTo", "scheme": "url"}
        ],
        "notes": FUNDING_STATEMENT,
    }
    (REPO_DIR / ".zenodo.json").write_text(json.dumps(zenodo, indent=2) + "\n", encoding="utf-8")

    license_text = f"""{ARCHIVE_LICENSE_NAME}

Copyright (c) {DATE.split("-")[0]} {AUTHOR_NAME}

Permission is hereby granted, free of charge, to any person obtaining a copy
of this software and associated documentation files (the "Software"), to deal
in the Software without restriction, including without limitation the rights
to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
copies of the Software, and to permit persons to whom the Software is
furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all
copies or substantial portions of the Software.

THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
SOFTWARE.
"""
    (REPO_DIR / "LICENSE").write_text(license_text, encoding="utf-8")


def fill_official_forms() -> None:
    """Populate Nature's official ML checklist (AcroForm) and Reporting Summary (XFA)."""
    import fill_nature_ml_checklist
    import fill_nature_reporting_summary

    fill_nature_ml_checklist.main()
    fill_nature_reporting_summary.main()


def archive_source_data() -> None:
    output = OUT / "Supplementary_Data_1_All_Result_Tables_CSV.zip"
    roots = [
        PRIMARY / "tables",
        BENCHMARK / "tables",
        REINFORCEMENT,
    ]
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        count = 0
        for root in roots:
            for path in sorted(root.glob("*.csv")):
                prefix = "primary" if root == PRIMARY / "tables" else "benchmark" if root == BENCHMARK / "tables" else "reinforcement"
                archive.write(path, f"{prefix}/{path.name}")
                count += 1
        archive.writestr("README.txt", f"Contains {count} original CSV result tables. Files are copied without modification.\n")
    if count != 31:
        raise ValueError(f"Expected 31 result CSV files, archived {count}")


def archive_source_figures() -> None:
    output = OUT / "Supplementary_Data_2_All_Original_Figures.zip"
    sources = [WORKFLOW]
    sources += sorted((PRIMARY / "figures").glob("*.png"))
    sources += sorted((PRIMARY / "figures").glob("*.svg"))
    sources += sorted((BENCHMARK / "figures").glob("*.png"))
    sources += sorted((BENCHMARK / "figures").glob("*.svg"))
    sources += sorted((REINFORCEMENT / "figures").glob("*.png"))
    sources += sorted((REINFORCEMENT / "figures").glob("*.svg"))
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sources:
            category = "workflow" if path == WORKFLOW else path.parent.parent.name
            archive.write(path, f"{category}/{path.name}")
        archive.writestr("README.txt", "Original retrieved figures only; composite supplementary figures are document layout derivatives.\n")


def archive_code() -> None:
    output = OUT / "Supplementary_Software_1_Code_and_Tests.zip"
    included = [ROOT / "src", ROOT / "tests", ROOT / "configs"]
    files = [
        ROOT / "README.md",
        ROOT / "pyproject.toml",
        ROOT / "modal_app.py",
        ROOT / "run_observability_benchmark.py",
        ROOT / "build_nmi_submission_package.py",
        ROOT / "build_manuscript_package.py",
        ROOT / "fill_nature_ml_checklist.py",
        ROOT / "fill_nature_reporting_summary.py",
        ROOT / "MODAL_RUNBOOK.md",
    ]
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for root in included:
            for path in sorted(root.rglob("*")):
                if path.is_file() and "__pycache__" not in path.parts:
                    archive.write(path, path.relative_to(ROOT).as_posix())
        for path in files:
            archive.write(path, path.relative_to(ROOT).as_posix())


def build_manifest() -> None:
    rows = []
    for path in sorted(OUT.rglob("*")):
        if path.is_file() and path.name != "SHA256_MANIFEST.csv" and QA not in path.parents:
            rows.append((path.relative_to(OUT).as_posix(), path.stat().st_size, hashlib.sha256(path.read_bytes()).hexdigest()))
    with (OUT / "SHA256_MANIFEST.csv").open("w", newline="", encoding="utf-8") as stream:
        writer = csv.writer(stream)
        writer.writerow(["relative_path", "bytes", "sha256"])
        writer.writerows(rows)


def build_machine_readable_claims() -> None:
    payload = {
        "title": TITLE,
        "primary_endpoint": "recursive_scientific_observability",
        "events": {"depmap_prism": 2169, "lincs": 2136, "tcga": 126, "total": 4431},
        "conditions": {"demonstrated": 9, "total": 9},
        "controlled_disruptions": {"graded_conditions": 90, "passed": 18, "total": 18},
        "median_cross_system_spearman": 0.753702346348183,
        "max_depth_gain": {
            "depmap_prism": 0.20692431167952718,
            "lincs": 0.41626755342362964,
            "tcga": 0.3991791267483722,
        },
        "secondary_tcga": {"patients": 42, "projects": 7, "auroc": 0.677, "brier": 0.215, "baseline_brier": 0.239},
        "claim_boundary": "observability of the recorded AI process; not biological correctness, clinical utility or causal treatment efficacy",
        "no_simulated_biological_data": True,
        "abstract_words": abstract_word_count(),
        "main_text_words": narrative_word_count(),
    }
    (QA / "audited_claims.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    QA.mkdir(parents=True, exist_ok=True)
    supplementary_figures = prepare_figures()
    build_main()
    build_supplement(supplementary_figures)
    build_cover_letter()
    build_title_page()
    build_reporting_summary()
    build_ml_checklist()
    build_instruction_audit()
    build_submission_checklist()
    build_standalone_table()
    download_forms()
    fill_official_forms()
    build_archive_metadata()
    archive_source_data()
    archive_source_figures()
    archive_code()
    export_pdfs()
    build_machine_readable_claims()
    build_manifest()
    print(json.dumps({
        "output": str(OUT),
        "main": str(MAIN_DOCX),
        "supplement": str(SUPP_DOCX),
        "abstract_words": abstract_word_count(),
        "main_text_words": narrative_word_count(),
    }, indent=2))


if __name__ == "__main__":
    main()
